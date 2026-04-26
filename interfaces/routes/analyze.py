from fastapi import APIRouter, Request, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
from typing import Optional
import os, time, logging, hashlib, io
from datetime import datetime, timezone, timedelta
import concurrent.futures as _cf
import httpx

from interfaces.config_runtime import _VAULT_SEMAPHORE_WS, _MODEL_LEIDA_A, _MODEL_LEIDA_B
from interfaces.models.schemas import ChatBeidni, ChatSvar, HeilsusvarModel
from interfaces.utils.helpers import _polish_fn_txt, _detect_filetype, _parse_docx, _parse_xlsx
from interfaces.utils.quota import FREE_QUOTA, _quota_tracker_doc, _beta_tracker, _er_beta_fras, _er_beta_ip, _promota_beta
from interfaces.utils.openrouter import _get_openrouter_balance, _wallet_preflight, _log_intent, SECURE_DOCS_DIR, MAX_PDF_SIZE
from interfaces.pipeline import _call_leid_a, _call_leid_b, _vault_system_prompt

logger = logging.getLogger("alvitur.web")

router = APIRouter()

@router.post("/api/analyze-document")
async def analyze_document(request: Request, file: Optional[UploadFile] = File(None), query: Optional[str] = Form(None)):
    """
    Sprint 21 — PDF Analyze Endpoint (Lag 1).
    Tekur við PDF skrá, les texta með PyMuPDF og búr til grófótta greiningu.
    Zero-Data: Skrá er eyð um leið og vinnslu ljúkur.
    Sprint 43b: query parameter wired from multipart form data.
    Sprint 43c: file optional — text-only queries fall back to direct LLM call.
    """

    # ── Sprint 43c: text-only fallback ──────────────────────────────
    _tier_hdr = request.headers.get("X-Alvitur-Tier", "general").lower().strip() if request else "general"
    _llm_model = _MODEL_LEIDA_B if _tier_hdr == "vault" else _MODEL_LEIDA_A
    if file is None or file.filename == "":
        if not query or not query.strip():
            return JSONResponse(status_code=422, content={
                "error_code": "empty_prompt",
                "detail": "Vinsamlegast skrifaðu fyrirspurn eða hladdu upp skjali.",
            })
        # Sprint 46 Phase 1b: quota check á text-only path (var vantar)
        import hashlib as _hl
        _master_key_txt = os.environ.get("ALVITUR_MASTER_KEY_HASH", "")
        _req_key_txt = request.headers.get("X-Master-Key", "") if request else ""
        _is_admin_txt = bool(_master_key_txt and _req_key_txt and _hl.sha256(_req_key_txt.encode()).hexdigest() == _master_key_txt)
        _client_ip_txt = (
            request.headers.get("CF-Connecting-IP")
            or request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
            or (request.client.host if request.client else "unknown")
        )
        # ── Sprint 64 A1: Beta check á text-only path (var vantar) ──
        # Speglarar logic í /api/chat (line ~3593) og analyze-document file-path (line ~3294)
        if _er_beta_fras(query or ""):
            _promota_beta(_client_ip_txt)
            logger.info(f"[BETA] {_client_ip_txt} promotaður í beta-tier (7d) via text-only")
        _is_beta_txt = _er_beta_ip(_client_ip_txt)
        # ── /Sprint 64 A1 beta check ──
        _log_intent("analyze-document/text-only", query, None, None, _tier_hdr)
        _quota_count_txt = _quota_tracker_doc.get(_client_ip_txt, 0) + 1
        if not _is_admin_txt and not _is_beta_txt:
            _quota_tracker_doc[_client_ip_txt] = _quota_count_txt
        if _quota_count_txt > FREE_QUOTA and not _is_admin_txt and not _is_beta_txt:
            return JSONResponse(status_code=403, content={
                "success": False,
                "error": "Ókeypis prufutími er liðinn.",
                "error_code": "quota_exceeded",
                "upgrade_required": True,
                "upgrade_url": "/askrift",
                "message": "Þú hefur nýtt þér 5 ókeypis beiðnir. Uppfærðu til að halda áfram.",
            })

        # Text-only path: send query directly to LLM (async — Sprint 46 Phase 1b)
        import os as _os
        import httpx as _httpx
        from datetime import datetime as _dt, timezone as _tz
        _key = _os.environ.get("OPENROUTER_API_KEY", "")
        _tier = _tier_hdr
        _summary = None
        _rag_txt = None  # tryggjum að nafnið sé til í lok return-blokks
        # 🟢 Sprint 62 Patch E v2: no-key → sovereign Leið B strax (skip OpenRouter entirely)
        if not _key and _tier_hdr != "vault":
            logger.warning("[ALVITUR] Sprint62E NO-KEY → sovereign fallback á Leið B")
            try:
                async with _VAULT_SEMAPHORE_WS:
                    _summary, _model_used, _usage = await _call_leid_b((query or "").strip())
            except Exception as _fe:
                logger.error(f"[ALVITUR] Sprint62E exc: {type(_fe).__name__}: {_fe}")
                _summary = None
            if _summary is None:
                return JSONResponse(status_code=503, content={
                    "error_code": "sovereign_unavailable",
                    "detail": "Staðbundin gervigreind er að ræsast. Reyndu eftir andartak."})
            logger.info(f"[ALVITUR] Sprint62E sovereign OK model={_model_used}")
            return JSONResponse(status_code=200, content={
                "success": True, "summary": _summary, "mode": "text-only",
                "pipeline_source": f"sovereign_nokey_{_model_used}", "model": _model_used})
        if _key:
            try:
                _now_str = _dt.now(_tz.utc).strftime("%Y-%m-%d %H:%M UTC")
                # Sprint 47: Domain classification + specialist prompt
                # Sprint 60c: None guard + honesty instruction
                from interfaces.specialist_prompts import get_specialist_prompt as _get_prompt
                from interfaces.skills.classify import ClassifySkill as _ClsSkill
                try:
                    _domain_txt = await _ClsSkill().run((query or "").strip() or "general", tier=_tier_hdr)
                except Exception:
                    _domain_txt = "general"
                _domain_txt = _domain_txt or "general"
                # s68-hotfix Part 3: text-only + tone-guide + thinking-suppress
                _honesty = (
                    "\n\nMikilvægt: Notandi hefur ekki hengt við skjal — þetta er almenn "
                    "spurning. Svaraðu út frá almennri þekkingu þinni á íslensku. "
                    "Ef þú veist ekki svarið með vissu, segðu það heiðarlega og bjóddu "
                    "framhaldsspurningu. Ekki búa til staðreyndir. "
                    "MIKILVÆGT UM FRAMSETNINGU: Ekki sýna hugsanaferli þitt — engar "
                    "<thinking>-tög, engin ‚bíð, ég þarf að“ setningar, engin internal-monologue. "
                    "Sendu aðeins hreint, prófessjónal svar beint. Engar emoji, "
                    "engin upphrópunarmerki, formlegt B2B-mál."
                )
                _system_prompt = _get_prompt(_domain_txt, _now_str) + _honesty
                # Sprint 70 Track D — RAG+ hook (text-only path)
                try:
                    from core.rag_orchestrator import retrieve_legal_context, build_rag_injection
                    _rag_txt = retrieve_legal_context(
                        query=(query or "").strip(),
                        intent_domain=_domain_txt,
                        tier=_tier,
                        tenant_id="system",
                    )
                    if _rag_txt.refusal:
                        from fastapi.responses import JSONResponse as _JR
                        return _JR(content={"success":True,"response":_rag_txt.refusal,
                            "pipeline_source":"rag_refusal_vault","domain":_domain_txt,
                            "zero_data":True,"found":True,"status":"ready_for_analysis",
                            "citations":[],"quota_warning":None,"tier":_tier})
                    if _rag_txt.used_retrieval:
                        _system_prompt = _system_prompt + build_rag_injection(_rag_txt.chunks)
                        _pipeline_source_txt = "rag_grounded_" + _tier
                    elif _rag_txt.fallback_to_gemini:
                        _system_prompt = _system_prompt + "[ATH: Engin lagatilvitnun fannst. Svaraðu varlega.]"
                        _pipeline_source_txt = "rag_fallback_general"
                except Exception as _rag_e2:
                    logger.warning("[RAG] text-only villa: %s", _rag_e2)
                logger.info(f"[ALVITUR] Sprint61 text-only tier={_tier} domain={_domain_txt}")
                _pipeline_source_txt = "unknown"
                if _tier == "vault":
                    from interfaces.config import VAULT_MAX_INPUT_TOKENS as _vmax
                    from interfaces.utils.helpers import _estimate_tokens
                    if _estimate_tokens(query or "") > _vmax:
                        return JSONResponse(status_code=413, content={
                            "error_code": "vault_input_too_large",
                            "detail": "Fyrirspurn er of stór fyrir Vault tier (max 8000 tokens). Styttu textann."})
                    async with _VAULT_SEMAPHORE_WS:
                        _summary, _model_used, _usage = await _call_leid_b(query.strip())
                    if _summary is None:
                        return JSONResponse(status_code=503, content={
                            "error_code": "vault_local_unavailable",
                            "detail": "Trúnaðarþjónusta tímabundið ekki tiltæk. Local AI module er að ræsast — reyndu aftur eftir 1 mínútu."})
                    _pipeline_source_txt = f"local_vllm_{_model_used}"
                    _in_tok = _usage.get("prompt_tokens", 0); _out_tok = _usage.get("completion_tokens", 0)
                    logger.info(f"[ALVITUR] leid_b sovereign tokens in={_in_tok} out={_out_tok}")
                else:
                    _summary, _model_used, _usage = await _call_leid_a(_system_prompt, query.strip())
                    if _summary is None:
                        logger.warning("[ALVITUR] Sprint62C text-only: Leið A None, reyni fallback á Leið B")
                        try:
                            async with _VAULT_SEMAPHORE_WS:
                                _summary, _model_used, _usage = await _call_leid_b(query.strip())
                        except Exception as _fe:
                            logger.error(f"[ALVITUR] Sprint62C fallback exc: {type(_fe).__name__}: {_fe}")
                            _summary = None
                        if _summary is None:
                            return JSONResponse(status_code=503, content={
                                "error_code": "both_pipelines_unavailable",
                                "detail": "Þjónusta tímabundið ekki aðgengileg. Reyndu eftir augnablik."})
                        _pipeline_source_txt = f"fallback_local_{_model_used}"
                        logger.info(f"[ALVITUR] Sprint62C fallback OK: model={_model_used}")
                    else:
                        _pipeline_source_txt = f"openrouter_{_model_used.split('/')[-1]}"
                    _in_tok = _usage.get("prompt_tokens", 0); _out_tok = _usage.get("completion_tokens", 0)
                    _cost = (_in_tok * 0.15 + _out_tok * 0.60) / 1_000_000 if "gpt-4o-mini" in _model_used else (_in_tok * 3.00 + _out_tok * 15.00) / 1_000_000 if "claude-sonnet" in _model_used else 0.0
                    logger.info("[ALVITUR] token_obs pipeline=text_query model=%s tier=%s in=%d out=%d cost=%.6f", _model_used, _tier, _in_tok, _out_tok, _cost)
                    try:
                        _summary = await _polish_fn_txt(_summary, _key)
                    except Exception as _pe:
                        logger.warning(f"[ALVITUR] polish failed (non-fatal): {_pe}")
            except Exception as _e:
                logger.error(f"[ALVITUR] text-only pipeline exc: {type(_e).__name__}: {_e}")
                _summary = None
                _domain_txt = "general"
                _pipeline_source_txt = "error"
        if _summary is None:
            # Sprint 62 Patch C: Final safety net — both pipelines down
            return JSONResponse(status_code=503, content={
                "error_code": "service_unavailable",
                "detail": "Greiningar þjónusta tímabundið ekki aðgengileg. Reyndu aftur eftir andartak.",
            })
        return JSONResponse(content={
            "success": True,
            "filename": None,
            "sidur": 0,
            "zero_data": True,
            "tier": _tier,
            "query": query.strip(),
            "response": _summary,
            "domain": _domain_txt,
            "citations": [],
            "found": True,
            "status": "ready_for_analysis",
            "pipeline_source": _pipeline_source_txt if "_pipeline_source_txt" in dir() else "unknown",
            "rag_metadata": _rag_txt.__dict__ if (_rag_txt is not None and hasattr(_rag_txt,"used_retrieval")) else {},
        })

    import fitz  # PyMuPDF

    # — 1. Staðfesta skrártegund —
    _ext = file.filename.lower().rsplit(".", 1)[-1] if "." in file.filename else ""
    if _ext not in ("pdf", "docx", "xlsx"):
        raise HTTPException(status_code=400, detail="Skráargerð ekki stuð. Sendu PDF, Word eða Excel.")

    # S2: Stream size check before full read
    _cl = None
    if hasattr(file, "headers") and file.headers:
        _cl = file.headers.get("content-length")
    if _cl and int(_cl) > MAX_PDF_SIZE:
        raise HTTPException(status_code=413,
            detail="Skjalið er yfir 20 MB — minnkaðu eða veldu annað.")
    efni = await file.read(MAX_PDF_SIZE + 1)
    if len(efni) > MAX_PDF_SIZE:
        raise HTTPException(status_code=413,
            detail="Skjalið er yfir 20 MB — minnkaðu eða veldu annað.")

    # K1/K2/S3: Filetype detection (magic bytes + extension)
    _filetype = _detect_filetype(efni, file.filename)


    # S4: Wallet circuit breaker
    _tier = request.headers.get("X-Alvitur-Tier", "general").lower().strip() if request else "general"
    _is_vault = (_tier == "vault")
    _wallet_preflight(is_vault=_is_vault)

    # PLG quota check — master key bypass
    _master_key = os.environ.get("ALVITUR_MASTER_KEY_HASH", "")
    _req_key = request.headers.get("X-Master-Key", "") if request else ""
    import hashlib as _hl
    _is_admin = bool(_master_key and _req_key and _hl.sha256(_req_key.encode()).hexdigest() == _master_key)
    _client_ip = (request.headers.get("CF-Connecting-IP") or request.headers.get("X-Forwarded-For", "").split(",")[0].strip() or (request.client.host if request.client else "unknown"))
    # Sprint 62 Patch A.1: define _is_beta early (was defined 240 lines later, causing NameError)
    try:
        _is_beta = _er_beta_ip(_client_ip)
    except Exception:
        _is_beta = False
    try:
        _fn = file.filename if file else None
        _fsize = getattr(file, "size", None) if file else None
    except Exception:
        _fn, _fsize = None, None
    _log_intent("analyze-document/file", query, _fn, _fsize,
                request.headers.get("X-Alvitur-Tier", "general") if request else "general")
    _quota_count = _quota_tracker_doc.get(_client_ip, 0) + 1
    if not _is_admin:
        _quota_tracker_doc[_client_ip] = _quota_count
    if _quota_count > FREE_QUOTA and not _is_admin and not _is_beta:
        return JSONResponse(status_code=403, content={
            "success": False,
            "error": "Ókeypis prufutími er liðinn.",
            "error_code": "quota_exceeded",
            "upgrade_required": True,
            "upgrade_url": "/askrift",
            "message": "Þú hefur nýtt þér 5 ókeypis beiðnir. Uppfærðu til að halda áfram.",
        })

    # PLG warning
    _quota_warning = None
    if _quota_count == 4 and not _is_admin:
        _quota_warning = "Þú hefur 1 beiðni eftir af 5 ókeypis beiðnum."
    elif _quota_count == 5 and not _is_admin:
        _quota_warning = "Þetta er þín síðasta ókeypis beiðni. Uppfærðu til að halda áfram."

    # S5+S6: Unified Parser (Office + PDF)
    # Office files (XLSX/DOCX)
    if _filetype in ('xlsx', 'docx'):
        import io
        try:
            if _filetype == 'xlsx':
                # Sprint 62 Patch B: Pandas pre-processor (Reiknivélar-Agent).
                # Returns markdown with real computed sums so LLM doesn't hallucinate.
                try:
                    from interfaces.excel_preprocessor import preprocess_excel as _prep_xlsx
                    heildartexti = _prep_xlsx(efni)
                    sidur = 1
                    logger.info(f"[ALVITUR] Sprint62b xlsx preprocessed via pandas, {len(heildartexti)} chars")
                except Exception as _xe:
                    logger.warning(f"[ALVITUR] pandas preprocess failed, fallback to openpyxl flat: {type(_xe).__name__}: {_xe}")
                    import openpyxl
                    wb = openpyxl.load_workbook(io.BytesIO(efni), read_only=True, data_only=True)
                    txt = []
                    for ws in wb.worksheets:
                        for row in ws.iter_rows(values_only=True):
                            clean = [str(c) if c is not None else "" for c in row]
                            line = " | ".join(v for v in clean if v.strip())
                            if line: txt.append(line)
                    heildartexti = "\n".join(txt)
                    sidur = len(wb.worksheets)
            elif _filetype == 'docx':
                from docx import Document
                doc = Document(io.BytesIO(efni))
                heildartexti = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                sidur = len(doc.paragraphs)
        except Exception as e:
            heildartexti = f"Villa við lestur: {str(e)}"
            sidur = 0
    else:
        # PDF Logic (Original)
        import concurrent.futures as _cf
        
        def _parse_pdf(data):
            parts, pages = [], 0
            with fitz.open(stream=data, filetype="pdf") as doc:
                pages = len(doc)
                for idx, pg in enumerate(doc):
                    t = pg.get_text().strip()
                    if t: parts.append(f"[Síða {idx+1}]\n{t}")
            return pages, parts

        texti_hlutar, sidur = [], 0
        if _is_vault:
            # VAULT: zero disk write
            with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                fut = ex.submit(_parse_pdf, efni)
                try:
                    sidur, texti_hlutar = fut.result(timeout=90)
                except _cf.TimeoutError:
                    raise HTTPException(status_code=504, detail="Vinnsla tók of langan tíma.")
        else:
            # GENERAL: disk write + cleanup
            skra_id = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S_%f")
            skra_slod = SECURE_DOCS_DIR / f"doc_{skra_id}.pdf"
            try:
                skra_slod.write_bytes(efni)
                with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                    fut = ex.submit(_parse_pdf, efni)
                    try:
                        sidur, texti_hlutar = fut.result(timeout=90)
                    except _cf.TimeoutError:
                        raise HTTPException(status_code=504, detail="Vinnsla tók of langan tíma.")
            finally:
                if skra_slod.exists(): skra_slod.unlink()
        
        heildartexti = "\n\n".join(texti_hlutar)

    if not heildartexti.strip():
        return JSONResponse(content={
            "success": False,
            "error": "Ekki tókst að lesa texta úr PDF (gæti verið skannað mynd).",
            "sidur": sidur,
        })

    # — 6. Skila grunnupplýsingum (LLM greining kemur í Lag 2) —
    # Sprint 62 Patch A: Only run fitz (PDF parser) for PDF files.
    # xlsx/docx already parsed above via openpyxl/python-docx.
    _parts = []
    if _filetype == 'pdf':
        import fitz as _fitz_inner
        with _fitz_inner.open(stream=efni, filetype="pdf") as _doc:
            sidur = len(_doc)
            for _pg in _doc:
                _t = _pg.get_text().strip()
                if _t: _parts.append(_t)
    # Default result fallback for all filetypes
    result = {"response": heildartexti if "heildartexti" in dir() else "", "citations": [], "found": bool(heildartexti.strip()) if "heildartexti" in dir() else False}

    # Build result dict for PDF fallback
    if _filetype == 'pdf':
        result = {"response": " ".join(_parts[:3000]), "citations": [], "found": bool(_parts)}
    import os as _os
    _tier = request.headers.get("X-Alvitur-Tier", "general").lower()
    _summary = None
    _pipeline_source_doc = "unknown"
    _domain_doc = "general"
    _rag_doc_meta = {"used_retrieval": False, "chunks_count": 0, "top_score": 0.0, "low_confidence": False, "source_laws": []}
    if _tier in ("general", "vault"):
        try:
            # 🟢 Sprint 62 Patch F: no-key → sovereign Leið B fyrir file-upload
            _key_preflight = _os.environ.get("OPENROUTER_API_KEY", "")
            if not _key_preflight and _tier != "vault":
                logger.warning("[ALVITUR] Sprint62F file-upload NO-KEY → sovereign Leið B")
                _msg_fallback = f"SPURNING: {(query or '').strip() or 'Greindu þetta skjal.'}\n\nSKJAL:\n{heildartexti[:30000]}"
                try:
                    async with _VAULT_SEMAPHORE_WS:
                        _summary, _model_used, _usage = await _call_leid_b(_msg_fallback)
                except Exception as _fe:
                    logger.error(f"[ALVITUR] Sprint62F exc: {type(_fe).__name__}: {_fe}")
                    _summary = None
                if _summary is not None:
                    _pipeline_source_doc = f"sovereign_nokey_file_{_model_used}"
                    logger.info(f"[ALVITUR] Sprint62F sovereign OK model={_model_used}")
                    # 🟢 Sprint 62 Patch I: EARLY RETURN — ekki keyra leið A aftur!
                    _in_tok = _usage.get("prompt_tokens", 0); _out_tok = _usage.get("completion_tokens", 0)
                    logger.info(f"[ALVITUR] Sprint62I early-return (skip leid_a) in={_in_tok} out={_out_tok}")
                    return JSONResponse(status_code=200, content={
                        "success": True,
                        "filename": _filename if "_filename" in dir() else "",
                        "sidur": sidur if "sidur" in dir() else 1,
                        "zero_data": True,
                        "tier": _tier,
                        "query": (query or "").strip(),
                        "response": _summary,
                        "domain": "general",
                        "citations": [],
                        "found": True,
                        "status": "ready_for_analysis",
                        "quota_warning": None,
                        "pipeline_source": _pipeline_source_doc,
                    })
            _key = _os.environ.get("OPENROUTER_API_KEY", "")
            _context_limit = 60000 if _tier == "vault" else 30000
            _full = heildartexti[:_context_limit]
            if _full:
                if query and query.strip():
                    truncation_note = " [Skjal styttur vegna stærðar]" if len(heildartexti) > _context_limit else ""
                    _msg = f"""SPURNING: {query.strip()}

REGLUR:
1. Svaraðu BEINT spurningunni hér að ofan á íslensku.
2. Notaðu skjalið AÐEINS sem heimild til að styðja við svarið.
3. EKKI endurskrifa skjalið, EKKI draga það saman nema beðið sé um það sérstaklega.
4. Ef upplýsingar vantar í skjalinu, segðu það hreint út.

SKJAL{truncation_note}:
{_full}"""
                else:
                    _msg = f"""Greindu þetta skjal stuttlega á íslensku (max 3 setningar).
Áhersla á meginþætti og ályktanir.

SKJAL:
{_full[:3000]}"""
                # Sprint 61: Leid A/B sovereign separation
                if _tier == "vault":
                    from interfaces.config import VAULT_MAX_INPUT_TOKENS as _vmax
                    if _estimate_tokens(_msg) > _vmax:
                        return JSONResponse(status_code=413, content={
                            "error_code": "vault_input_too_large",
                            "detail": f"Skjal er of stórt fyrir Vault tier (max {_vmax} tokens). Skiptu í smærri hluta eða notaðu Almenna greiningu."})
                    logger.info(f"[ALVITUR] Sprint61 analyze_doc tier=vault calling leid_b")
                    async with _VAULT_SEMAPHORE_WS:
                        _summary, _model_used, _usage = await _call_leid_b(_msg)
                    if _summary is None:
                        return JSONResponse(status_code=503, content={
                            "error_code": "vault_local_unavailable",
                            "detail": "Trúnaðarþjónusta tímabundið ekki tiltæk. Local AI module er að ræsast — reyndu aftur eftir 1 mínútu."})
                    _pipeline_source_doc = f"local_vllm_{_model_used}"
                    _in_tok = _usage.get("prompt_tokens", 0); _out_tok = _usage.get("completion_tokens", 0)
                    logger.info(f"[ALVITUR] leid_b analyze_doc in={_in_tok} out={_out_tok}")
                    _domain_doc = "vault"
                else:
                    from datetime import datetime as _dt, timezone as _tz
                    from interfaces.specialist_prompts import classify as _classify, get_specialist_prompt as _get_prompt
                    _classify_text = (query.strip() if query and query.strip() else heildartexti[:500])
                    try:
                        _domain_doc = await _classify(_classify_text or "general")
                    except Exception:
                        _domain_doc = "general"
                    _domain_doc = _domain_doc or "general"
                    _now_str = _dt.now(_tz.utc).strftime("%Y-%m-%d %H:%M UTC")
                    # s68-hotfix: conditional honesty — document vs text-only mode
                    _has_document = bool(file and getattr(file, "filename", "") and file.filename.strip())
                    if _has_document:
                        _honesty_doc = (
                            "\n\nMikilvægt: Ef þú finnur ekki nógu nákvæmar upplýsingar í skjalinu "
                            "til að svara spurningunni, segjum notandanum það beint og bjóðum upp á "
                            "framhaldsspurningu. Búðu ALDREI til upplýsingar sem eru ekki í skjalinu."
                        )
                    else:
                        # s68-hotfix Part 3: tone-guide + thinking-suppress
                        _honesty_doc = (
                            "\n\nMikilvægt: Notandi hefur ekki hengt við skjal — þetta er almenn "
                            "spurning. Svaraðu út frá almennri þekkingu þinni á íslensku. "
                            "Ef þú veist ekki svarið með vissu, segðu það heiðarlega og bjóddu "
                            "framhaldsspurningu. Ekki búa til staðreyndir. "
                            "MIKILVÆGT UM FRAMSETNINGU: Ekki sýna hugsanaferli þitt — engar "
                            "<thinking>-tög, engin ‚bíð, ég þarf að“ setningar, engin internal-monologue. "
                            "Sendu aðeins hreint, prófessjónal svar beint. Engar emoji, "
                            "engin upphrópunarmerki, formlegt B2B-mál."
                        )
                    _system_prompt = _get_prompt(_domain_doc, _now_str) + _honesty_doc

                    # Sprint 70 Track D — RAG+ hook
                    print("DEBUG: RAG hook starting", flush=True)
                    try:
                        from core.rag_orchestrator import retrieve_legal_context, build_rag_injection
                        _rag = retrieve_legal_context(
                            query=((query or "").strip() or _msg or ""),
                            intent_domain=_domain_doc,
                            tier="general",
                            tenant_id="system",
                        )
                        if _rag.refusal:
                            return JSONResponse(content={
                                "success": True, "response": _rag.refusal,
                                "pipeline_source": "rag_refusal_vault",
                                "domain": _domain_doc, "zero_data": True,
                                "found": True, "status": "ready_for_analysis",
                                "citations": [], "quota_warning": None,
                            })
                        if _rag.used_retrieval:
                            _rag_injection = build_rag_injection(_rag.chunks)
                            _system_prompt = _system_prompt + _rag_injection
                            _pipeline_source_doc = f"rag_grounded_general"
                            _rag_doc_meta = {"used_retrieval": True, "chunks_count": len(_rag.chunks), "top_score": round(_rag.top_score,3), "low_confidence": _rag.top_score<0.65, "source_laws": []}
                            logger.info("[RAG] injected %d chunks into analyze_doc prompt", len(_rag.chunks))
                        elif _rag.fallback_to_gemini:
                            _system_prompt = _system_prompt + "[ATH: Engin lagatilvitnun fannst i corpus Alvitur. Svaraðu varlega.]"
                            _pipeline_source_doc = "rag_fallback_general"
                            _pipeline_source_doc = "rag_fallback_general"
                    except Exception as _rag_e:
                        logger.warning("[RAG] orchestrator villa (graceful degradation): %s", _rag_e)

                    _rag_doc_meta = {"used_retrieval": _rag.used_retrieval, "chunks_count": len(_rag.chunks), "top_score": round(_rag.top_score, 3), "low_confidence": _rag.top_score < 0.65, "source_laws": []} if "_rag" in dir() else {}
                    logger.info(f"[ALVITUR] Sprint61 analyze_doc tier=general calling leid_a domain={_domain_doc}")
                    _summary, _model_used, _usage = await _call_leid_a(_system_prompt, _msg)
                    if _summary is None:
                        # Sprint 62 Patch C: Leið A klikkaði → sovereign fallback á Leið B
                        logger.warning("[ALVITUR] Sprint62C analyze_doc: Leið A None, reyni fallback á Leið B")
                        try:
                            async with _VAULT_SEMAPHORE_WS:
                                _summary, _model_used, _usage = await _call_leid_b(_msg)
                        except Exception as _fe:
                            logger.error(f"[ALVITUR] Sprint62C analyze_doc fallback exc: {type(_fe).__name__}: {_fe}")
                            _summary = None
                        if _summary is None:
                            return JSONResponse(status_code=503, content={
                                "error_code": "both_pipelines_unavailable",
                                "detail": "Þjónusta tímabundið ekki aðgengileg. Reyndu eftir augnablik."})
                        _pipeline_source_doc = f"fallback_local_{_model_used}"
                        logger.info(f"[ALVITUR] Sprint62C analyze_doc fallback OK: model={_model_used}")
                    else:
                        _pipeline_source_doc = f"openrouter_{_model_used.split('/')[-1]}"
                    _in_tok = _usage.get("prompt_tokens", 0); _out_tok = _usage.get("completion_tokens", 0)
                    _cost = (_in_tok * 0.15 + _out_tok * 0.60) / 1_000_000 if "gpt-4o-mini" in _model_used else (_in_tok * 3.00 + _out_tok * 15.00) / 1_000_000 if "claude-sonnet" in _model_used else 0.0
                    logger.info("[ALVITUR] token_obs pipeline=analyze_doc model=%s tier=%s in=%d out=%d cost=%.6f", _model_used, _tier, _in_tok, _out_tok, _cost)
                    # Polish only for Leid A (never vault — sovereignty)
                    try:
                        # from interfaces.chat_routes import _polish as _polish_fn
                        _summary = await _polish_fn(_summary, _key)
                    except Exception as _pe:
                        logger.warning(f"[ALVITUR] polish failed (non-fatal): {_pe}")
        except Exception as _e:
            logger.error(f"[ALVITUR] analyze_doc pipeline exc: {type(_e).__name__}: {_e}")
            _summary = None
            _domain_doc = "general"
            _pipeline_source_doc = "error"
    response = {
        "success": True,
        "filename": file.filename,
        "sidur": sidur,
        "zero_data": True,
        "tier": _tier,
        "query": query or "",
        "response": _summary or result.get("response", ""),
        "domain": _domain_doc if "_domain_doc" in dir() else "general",
        "citations": result.get("citations", []),
        "found": result.get("found", False),
        "status": "ready_for_analysis",
        "quota_warning": _quota_warning,
        "rag_metadata": _rag_doc_meta,
        "pipeline_source": _pipeline_source_doc if "_pipeline_source_doc" in dir() else "unknown",
    }
    return JSONResponse(content=response)

