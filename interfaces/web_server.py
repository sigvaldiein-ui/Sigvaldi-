"""
Alvitur.is — Enterprise AI Landing Page
FastAPI server serving a one-page B2B SaaS funnel for Icelandic AI document analysis.

Sprint 16 breytingar:
  A) Web Chat Fasi 1 samþætt á /minarsidur (mock svör, engin V5 tenging)
  B) Sölutexti uppfærður — breiðari markhópur (lögfræðistofur, fjármálafyrirtæki, ráðgjafar)
  C) Dæmisögur ("Notendur segja") bætt við forsíðuna
  D) SecurityHeadersMiddleware bætt við (HSTS, CSP, X-Frame-Options, o.fl.)

Sprint 17 hotfix — Öryggi + UX:
  1) Netfang skipt: sigvaldi@fjarmalin.is / sigvaldiein@gmail.com → info@alvitur.is
  2) Beta upplýsingakassi bætt við forsíðuna undir hero (Telegram hlekkur)
  3) Session einangrun staðfest á /api/chat: engin shared state milli notenda

Sprint 18 — Hreinsa vefspjall / Web chat cleanup:
  1) Fjarlægður 'Mínar Síður' hlekkur úr nav
  2) Fjarlægður 'Opna á Telegram' hnappur úr nav (einungis Alvitur .is merki og valmynd eftir)
  3) Skipt út tveimur hnöppum undir beta kassa fyrir einn: ''
  4) /minarsidur route og build_minarsidur_page() commentuð út (geymd til sögunnar)
  5) Öll /minarsidur tenglar á forsíðu skipt út fyrir Telegram / mailto: hlekki
  6) Static mount (/static) commentuð út — var eingöngu notuð fyrir /minarsidur
  7) /api/chat route haldið (virkt — til framtíðarnotkunar á Fasa 2)

Sprint 19 — UI Hotfix (Aðal áætlun, innleiðsla Per):
  1) Telegram hraðhnappur aftur í navbar (sýnilegur á desktop + mobile)
  2) Hero padding þétt — meira efni above-the-fold
  3) Status pilla (hero-badge) smækkuð — minni visual weight
  4) Beta kassi fágaður — mildari litir, premium dökkt útlit
  5) Mobile-first fókus — Telegram hnappur alltaf sýnilegur, þéttara hero

Sprint 20 — V5.1 B2B Evidence Engine (Aðal áætlun, innleiðsla Per):
  1) H1 breytt í 'Alvitur – Við greinum gögnin sem skipta máli.'
  2) Undirfyrirsögn: Zero-Data, B2B markhópur
  3) Aðalhnappur: 'Óska eftir prufuaðgangi' → mailto:info@alvitur.is
  4) Drag-and-Drop PDF svæði bætt við (óvirkt viðmót, Evidence Engine tónn)
  5) Meta title og description uppfært
"""

import asyncio
import logging
import os

# Sprint 63 Track A: Load .env BEFORE any os.environ.get() calls
# Sprint 63 Track B1: env-aware loading — .env.dev if ALVITUR_ENV=dev
from dotenv import load_dotenv
import os as _os_init
_env_flag = _os_init.environ.get("ALVITUR_ENV", "prod")
_env_file = "/workspace/.env.dev" if _env_flag == "dev" else "/workspace/.env"
load_dotenv(_env_file)

# Sprint 63 Track A5: track server uptime for /api/diagnostics
import time as _time_init
_SERVER_START_TIME = _time_init.time()

# Sprint 63 Track A6: polish stub (restore removed-in-refactor function)
async def _polish_fn_txt(*args, **kwargs) -> str:
    """Async no-op polish stub — Sprint 63 Track A6.2.
    Caller notar 'await' svo þetta er async. Skilar textanum óbreyttum.
    TODO: bæta við raunverulegri íslenskri málfræðipúlisingu síðar.
    """
    text = args[0] if args else kwargs.get("text", "")
    return text if isinstance(text, str) else str(text)
import sys as _sys
_repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _repo_root not in _sys.path:
    _sys.path.insert(0, _repo_root)
import random
import time
from collections import defaultdict
from datetime import datetime, timezone
from typing import Optional

from pathlib import Path
from fastapi import FastAPI, HTTPException, Request, status, UploadFile, File, Form
from fastapi.exceptions import RequestValidationError
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.middleware.gzip import GZipMiddleware
from pydantic import BaseModel, Field, validator
# Sprint 64 B2-V2: Intent Gateway observability (lazy import, never raises)
try:
    import sys as _sys, os as _os
    _here = _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__)))
    if _here not in _sys.path:
        _sys.path.insert(0, _here)
    from core.intent_gateway import classify_intent as _classify_intent
    _INTENT_AVAILABLE = True
except Exception:
    _classify_intent = None
    _INTENT_AVAILABLE = False, validator
from starlette.middleware.base import BaseHTTPMiddleware
import uvicorn

# ─── Lög stillingar ───────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
import asyncio as _aio_ws
from interfaces.static_content.html_pages import (
    SHARED_HEAD, SHARED_STYLES,
    SVG_LOGO, SVG_LOGO_CHAT, CHECK_SVG, ACCENT_CHECK, NAV_SCRIPT, HTML_PAGE,
)
from interfaces.pipeline import _call_leid_a, _call_leid_b, _vault_system_prompt
_VAULT_SEMAPHORE_WS = _aio_ws.Semaphore(1)
logger = logging.getLogger("alvitur.web")

# Sprint 54: pipeline_adapter skilar alltaf None — fallback path active

# ─── Gæðatak (Rate Limiting) ──────────────────────────────────────────────────
# Hámark 20 beiðnir á mínútu á hvert IP
GÆÐATAK_HÁMARK = int(os.environ.get("RATE_LIMIT", 20))
GÆÐATAK_GLUGGI = 60  # sekúndur

# Sprint 53b: Módel stillingar frá config.py — aldrei hardcoded
from interfaces.config import MODEL_LEIDA_A as _MODEL_LEIDA_A, MODEL_LEIDA_B as _MODEL_LEIDA_B

# Geymir: {ip_address: [tímastimplar]} fyrir síðustu 60 sekúndur
_gæðatak_minni: dict[str, list[float]] = defaultdict(list)


def athuga_gæðatak(ip: str) -> bool:
    """
    Sannprófar hvort IP tala sé yfir gæðataki.
    Skilar True ef beiðni er leyfð, False ef yfir mörk.
    Notar gluggatíma (sliding window) — einfaldasta útfærslan án Redis.
    """
    núna = time.time()
    _gæðatak_minni[ip] = [t for t in _gæðatak_minni[ip] if núna - t < GÆÐATAK_GLUGGI]
    if len(_gæðatak_minni[ip]) >= GÆÐATAK_HÁMARK:
        return False
    _gæðatak_minni[ip].append(núna)
    return True


def sækja_ip(request: Request) -> str:
    """Sækir raunverulegt IP-tölu notanda, tekur tillit til proxy."""
    for haus in ("cf-connecting-ip", "x-real-ip"):
        if ip := request.headers.get(haus):
            return ip.split(",")[0].strip()
    return request.client.host if request.client else "0.0.0.0"


# ─── D) Öryggishausar Middleware ──────────────────────────────────────────────

class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    """
    Bætir við öryggishausum í hverja HTTP svörun.

    Hausar sem eru settir:
      - Strict-Transport-Security (HSTS)     : HTTPS alltaf, 1 ár
      - Content-Security-Policy              : XSS vörn
      - X-Frame-Options                      : Kemur í veg fyrir clickjacking
      - X-Content-Type-Options               : Kemur í veg fyrir MIME sniffing
      - Referrer-Policy                      : Takmarkar referrer upplýsingar
      - Permissions-Policy                   : Takmarkar vafra API aðgang
    """

    async def dispatch(self, request: Request, call_next):
        svar = await call_next(request)

        # HSTS — krefjast HTTPS í 1 ár með undirlénum
        svar.headers["Strict-Transport-Security"] = (
            "max-age=31536000; includeSubDomains"
        )

        # CSP — takmarkar hvaðan efni má sækja
        # Athugasemd: unsafe-inline er í stílar vegna innilagðra stíla (Fasi 1).
        # Í Fasa 2 færum við CSS í utanaðkomandi skrá og fjarlægjum unsafe-inline.
        svar.headers["Content-Security-Policy"] = (
            "default-src 'self'; "
            "script-src 'self' 'unsafe-inline'; "
            "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com https://api.fontshare.com; "
            "font-src 'self' https://fonts.gstatic.com https://api.fontshare.com data:; "
            "img-src 'self' data:; "
            "connect-src 'self'; "
            "frame-ancestors 'none'; "
            "base-uri 'self';"
        )

        # X-Frame-Options — kemur í veg fyrir að síðan sé sett í iframe (clickjacking)
        svar.headers["X-Frame-Options"] = "DENY"

        # X-Content-Type-Options — kemur í veg fyrir MIME tegundaspá
        svar.headers["X-Content-Type-Options"] = "nosniff"

        # Referrer-Policy — takmarkar hvaðan referrer er sent
        svar.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"

        # Permissions-Policy — takmarkar vafra API aðgang
        svar.headers["Permissions-Policy"] = (
            "geolocation=(), microphone=(), camera=(), payment=()"
        )

        return svar


# ─── FastAPI forrit ────────────────────────────────────────────────────────────

app = FastAPI(
    title="Alvitur Enterprise AI",
    docs_url=None,   # Slökkva á Swagger UI í framleiðslu
    redoc_url=None,  # Slökkva á ReDoc í framleiðslu
)

# Sprint 43b: Custom 422 handler — add error_code for frontend compatibility
@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    # Check if the error is specifically about a missing file field
    error_code = "validation_error"
    for err in exc.errors():
        if err.get("loc") and "file" in err.get("loc", []):
            error_code = "no_file"
            break
    return JSONResponse(
        status_code=422,
        content={
            "error": "validation_error",
            "error_code": error_code,
            "detail": exc.errors(),
        },
    )

# Gzip þjöppun (sparar bandbreidd)
app.add_middleware(GZipMiddleware, minimum_size=1000)

# D) Öryggishausar á allar beiðnir
app.add_middleware(SecurityHeadersMiddleware)

# Sprint 18: Static mount commentuð út — var eingöngu notuð fyrir /minarsidur vefspjall
from fastapi.staticfiles import StaticFiles
# import os as _os
_static_dir = "/workspace/mimir_net/interfaces/static"
# _os.makedirs(_static_dir, exist_ok=True)
app.mount("/static", StaticFiles(directory=_static_dir), name="static")


# ─── A) Gagnalíkön fyrir Web Chat API ────────────────────────────────────────

class ChatBeidni(BaseModel):
    """Inntak fyrir /api/chat."""
    message: str = Field(
        ...,
        min_length=1,
        max_length=4000,
        description="Skilaboð frá notanda",
    )
    user_id: str = Field(
        default="anonymous",
        max_length=64,
        description="Notandaauðkenni (tímabundið, skipt út á Fasa 3)",
    )

    @validator("message")
    def hreinsa_skilabus(cls, v: str) -> str:
        """Þrífa og sannreyna skilaboð."""
        hrein = v.strip()
        if not hrein:
            raise ValueError("Skilaboð má ekki vera tómt")
        return hrein


class ChatSvar(BaseModel):
    """Úttak frá /api/chat."""
    response: str
    timestamp: str
    # [FASI-2] Bæta við: model_used, tokens_used, search_used
    # [FASI-3] Bæta við: query_remaining


class HeilsusvarModel(BaseModel):
    """Úttak frá /api/health."""
    status: str
    version: str
    timestamp: str
    fasi: str


# ─── A) Mock svörþjónusta ─────────────────────────────────────────────────────
# [FASI-2]: Skipta þessum mock svörum út fyrir MimirCoreV5.process_message()

# Íslenskt velkomnarboð
VELKOMIN_BOÐ = (
    "Góðan daginn! Ég er **Alvitur**, íslenskur gervigreindaraðstoðarmaður "
    "sérhæfður í skjalagreiningu og lögfræðilegri rannsókn.\n\n"
    "Ég get aðstoðað þig með:\n"
    "- Greiningu á **ársreikningum** og fjárhagslegum skjölum\n"
    "- Leit og samantekt úr **íslenskum lögum og reglugerðum**\n"
    "- Samanburð á **samningum** og lagalegum gögnum\n"
    "- **Minnisblöð** og skýrslugerð á sekúndum\n\n"
    "Hvað get ég gert fyrir þig í dag?"
)

# ─── Shared Styles & Components ───────────────────────────────────────────────


# ─── SVG Logo ─────────────────────────────────────────────────────────────────

# ─── Check SVG ────────────────────────────────────────────────────────────────

# ─── Nav Script ───────────────────────────────────────────────────────────────

# ─── Sprint 44: Frontend Redesign — Light theme ──────────────────────────────



# ─── A) Web Chat Fasi 1 — /minarsidur [FJARLÆGT Sprint 18] ────────────────────
# Samþætt úr web_chat_fasi1.py — fullbúið UI með mock svörum
# Sprint 18: build_minarsidur_page() og /minarsidur route commentuð út — vefspjall hreinað
# [FASI-2]: Tengja við MimirCoreV5.process_message() í stað mock falls
# def build_minarsidur_page() -> str:
#     """
#     Smíðar HTML síðuna fyrir vefspjallið á /minarsidur.
#     Notar sömu CSS breytur og forsíðan (dark mode Enterprise hönnun).
#     Fasi 1: Mock svör, engin JWT auðkenning, engin V5 tenging.
#     """
#     return """<!DOCTYPE html>
# <html lang="is">
# <head>
#   <meta charset="UTF-8">
#   <meta name="viewport" content="width=device-width, initial-scale=1.0">
#   <meta name="description" content="Alvitur — Íslenskur AI aðstoðarmaður fyrir skjalagreiningu">
#   <meta name="robots" content="noindex, nofollow">
#   <title>Alvitur — Lokað Gagnaherbergi</title>
#   {SHARED_HEAD}
#   {SHARED_STYLES}
#   <style>
#     /* Yfirskrifa body til fullskjás í spjalli */
#     body {{
#       overflow: hidden;
#       min-height: 100dvh;
#     }}
#   </style>
# </head>
# <body>
# 
#   <!-- Aðgengistengill -->
#   <a href="#spjall-inntak" class="skip-link">Fara beint í inntak</a>
# 
#   <!-- ─── HAUS ─────────────────────────────── -->
#   <header class="chat-haus" role="banner">
#     <a href="/" class="chat-til-baka" aria-label="Til baka á alvitur.is forsíðu">
#       <svg width="14" height="14" viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="2" aria-hidden="true">
#         <path d="M10 12L6 8l4-4"/>
#       </svg>
#       <span>Til baka</span>
#     </a>
#     <div class="chat-aðskilir" aria-hidden="true"></div>
#     <div class="chat-logo" aria-label="Alvitur">
#       {SVG_LOGO_CHAT}
#       Alvitur
#     </div>
#     <div class="chat-merki" role="status" aria-label="Þjónusta virk">
#       <div class="chat-merki-dot" aria-hidden="true"></div>
#       Lokað Gagnaherbergi
#     </div>
#   </header>
# 
#   <!-- ─── AÐALSVÆÐI ────────────────────────── -->
#   <div class="chat-skipulag">
# 
#     <!-- Spjallsaga -->
#     <main class="chat-saga" id="spjall-saga" role="log" aria-live="polite" aria-label="Spjallsaga" tabindex="0">
#       <div class="chat-innri" id="spjall-innri">
# 
#         <!-- Velkomnarboð (birtist við hleðslu) -->
#         <div class="boð-lína ai-boð" id="velkomin-boð">
#           <div class="boð-merki ai-merki" aria-hidden="true">
#             <svg width="16" height="16" viewBox="0 0 32 32" fill="none">
#               <path d="M16 2L28.7 9.5V24.5L16 32L3.3 24.5V9.5L16 2Z" fill="rgba(99,102,241,0.3)" stroke="#818CF8" stroke-width="1.5"/>
#               <path d="M11 22L16 10L21 22M13 19H19" stroke="#818CF8" stroke-width="1.75" stroke-linecap="round" stroke-linejoin="round"/>
#             </svg>
#           </div>
#           <div>
#             <div class="boð-bubble">
#               <div class="boð-texti" id="velkomin-texti"></div>
#             </div>
#             <span class="boð-tími" id="velkomin-timi" aria-label="Tímastimpill"></span>
#           </div>
#         </div>
# 
#       </div>
#     </main>
# 
#     <!-- Inntakssvæðið -->
#     <form id="chat-form" class="inntak-sviði" role="region" aria-label="Skilaboðainntak" onsubmit="return false;">
#       <div class="inntak-innri">
#         <div class="inntak-rowr">
#           <!-- Skráarsendir (óvirkur í Fasa 1) -->
#           <button
#             class="inntak-skrá-hnappurinn"
#             disabled
#             aria-label="Hlaða upp skjali (kemur fljótlega)"
#             title="Skráarsendir kemur í Fasa 4"
#           >
#             <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" aria-hidden="true">
#               <path d="M21.44 11.05l-9.19 9.19a6 6 0 01-8.49-8.49l9.19-9.19a4 4 0 015.66 5.66l-9.2 9.19a2 2 0 01-2.83-2.83l8.49-8.48"/>
#             </svg>
#           </button>
# 
#           <!-- Textasvæðið -->
#           <div class="inntak-þekja">
#             <textarea
#               id="spjall-inntak"
#               class="inntak-textarea"
#               placeholder="Skrifaðu spurninguna þína hér..."
#               rows="1"
#               aria-label="Spjallinntak"
#               aria-multiline="true"
#             ></textarea>
#           </div>
# 
#           <!-- Sendahnappur -->
#           <button
#             id="senda-hnappurinn"
#             class="inntak-senda-hnappurinn"
#             aria-label="Senda skilaboð"
#             onclick="if(typeof handleSend==='function')handleSend();"
#           >
#             <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" aria-hidden="true">
#               <path d="M22 2L11 13M22 2L15 22 11 13 2 9l20-7z" stroke-linecap="round" stroke-linejoin="round"/>
#             </svg>
#           </button>
#         </div>
#         <p class="inntak-leiðbeiningar">
#           <kbd>Enter</kbd> sendir · <kbd>Shift+Enter</kbd> = nýlína ·
#           <span style="color: var(--accent-light)">Fasi 1 — Mock svör</span>
#         </p>
#       </div>
#     </form>
# 
#   </div>
# 
#   <!-- ─── A) VEFSPJALL JAVASCRIPT ──────────────── -->
#   <script data-version="99">
#   // ================================================================
#   // Alvitur Vefspjall — Fasi 1 (Mock svör)
#   // Sprint 16.4 — samþætt úr web_chat_fasi1.py
#   //
#   // [FASI-2]: Skipta /api/chat kalli út fyrir MimirCoreV5 tengingu
#   // [FASI-3]: Bæta við JWT auðkenningu og lotustjórnun
#   // [FASI-4]: SSE streyming, PDF upphleðsla
#   // ================================================================
# 
#   // ─── Hjálparföll ──────────────────────────────────────────────
#   function currentTime() {{
#     return new Date().toLocaleTimeString('is-IS', {{
#       hour: '2-digit',
#       minute: '2-digit'
#     }});
#   }}
# 
#   // Einföldur Markdown þáttur (Fasi 1 — regex)
#   // [FASI-2]: Skipta út fyrir marked.js CDN
#   function parseMarkdown(texti) {{
#     // Þrífa og varða HTML
#     let html = texti
#       .replace(/&/g, '&amp;')
#       .replace(/</g, '&lt;')
#       .replace(/>/g, '&gt;');
# 
#     // Markdown → HTML umbreytingar (röð skiptir máli)
#     html = html
#       // Kóðablokk (```...```)
#       .replace(/```([\\s\\S]*?)```/g, '<pre><code>$1</code></pre>')
#       // Inline kóði (`...`)
#       .replace(/`([^`]+)`/g, '<code>$1</code>')
#       // Feitletrað (**...**)
#       .replace(/\\*\\*([^*]+)\\*\\*/g, '<strong>$1</strong>')
#       // Skáletrað (*...*)
#       .replace(/\\*([^*]+)\\*/g, '<em>$1</em>')
#       // Markdown tafla — einföldun
#       .replace(/\\|(.+)\\|/g, (match) => {{
#         const frumur = match.split('|').filter(s => s.trim() !== '');
#         const isSeparator = frumur.every(s => /^[-:\\s]+$/.test(s));
#         if (isSeparator) return '';
#         const td = frumur.map(s => `<td>${{s.trim()}}</td>`).join('');
#         return `<tr>${{td}}</tr>`;
#       }})
#       // Pakka töflurowm
#       .replace(/(<tr>.*<\\/tr>)+/gs, match => {{
#         const rowr = match.match(/<tr>.*?<\\/tr>/gs) || [];
#         if (rowr.length === 0) return match;
#         const haus = rowr[0].replace(/<td>/g, '<th>').replace(/<\\/td>/g, '</th>');
#         const tbody = rowr.slice(1).join('');
#         return `<table><thead>${{haus}}</thead><tbody>${{tbody}}</tbody></table>`;
#       }})
#       // Listi (- atriði)
#       .replace(/^- (.+)$/gm, '<li>$1</li>')
#       .replace(/(<li>.*<\\/li>\\n?)+/gs, match => `<ul>${{match}}</ul>`)
#       // Númeraður listi
#       .replace(/^\\d+\\.(.+)$/gm, '<li>$1</li>')
#       // Tenglar [texti](URL)
#       .replace(/\\[([^\\]]+)\\]\\(([^)]+)\\)/g, '<a href="$2" target="_blank" rel="noopener noreferrer">$1</a>')
#       // Línufall → efnisgrein
#       .replace(/\n\n/g, '</p><p>')
#       // Einföldur rowfall
#       .replace(/\n/g, '<br>');
# 
#     // Pakka í efnisgrein ef ekki þegar pakkað
#     if (!html.startsWith('<')) {{
#       html = `<p>${{html}}</p>`;
#     }}
# 
#     return html;
#   }}
# 
#   // ─── Búa til boðarow ──────────────────────────────────────────
#   function createAiMsg(texti, er_hlaðning = false) {{
#     const row = document.createElement('div');
#     row.className = 'boð-lína ai-boð';
# 
#     const merki = document.createElement('div');
#     merki.className = 'boð-merki ai-merki';
#     merki.setAttribute('aria-hidden', 'true');
#     merki.innerHTML = `<svg width="16" height="16" viewBox="0 0 32 32" fill="none">
#       <path d="M16 2L28.7 9.5V24.5L16 32L3.3 24.5V9.5L16 2Z" fill="rgba(99,102,241,0.3)" stroke="#818CF8" stroke-width="1.5"/>
#       <path d="M11 22L16 10L21 22M13 19H19" stroke="#818CF8" stroke-width="1.75" stroke-linecap="round" stroke-linejoin="round"/>
#     </svg>`;
# 
#     const efni = document.createElement('div');
#     const bubble = document.createElement('div');
#     bubble.className = 'boð-bubble';
# 
#     const msgText = document.createElement('div');
#     msgText.className = 'boð-texti';
# 
#     if (er_hlaðning) {{
#       // Þrír boppandi punktar meðan beðið er
#       msgText.innerHTML = `<div class="hladning-rowr" aria-label="Alvitur er að hugsa...">
#         <div class="hladning-dot"></div>
#         <div class="hladning-dot"></div>
#         <div class="hladning-dot"></div>
#       </div>`;
#     }} else {{
#       msgText.innerHTML = parseMarkdown(texti);
#     }}
# 
#     bubble.appendChild(msgText);
# 
#     const timi = document.createElement('span');
#     timi.className = 'boð-tími';
#     timi.textContent = currentTime();
#     timi.setAttribute('aria-label', `Sent kl. ${{currentTime()}}`);
# 
#     efni.appendChild(bubble);
#     efni.appendChild(timi);
#     row.appendChild(merki);
#     row.appendChild(efni);
# 
#     return {{ row, msgText }};
#   }}
# 
#   function createUserMsg(texti) {{
#     const row = document.createElement('div');
#     row.className = 'boð-lína notenda-boð';
# 
#     const efni = document.createElement('div');
#     const bubble = document.createElement('div');
#     bubble.className = 'boð-bubble';
#     bubble.textContent = texti;
# 
#     const timi = document.createElement('span');
#     timi.className = 'boð-tími';
#     timi.textContent = currentTime();
# 
#     efni.appendChild(bubble);
#     efni.appendChild(timi);
#     row.appendChild(efni);
# 
#     return row;
#   }}
# 
#   // ─── Skruna neðst ──────────────────────────────────────────────
#   function scrollBottom() {{
#     const saga = document.getElementById('spjall-saga');
#     saga.scrollTop = saga.scrollHeight;
#   }}
# 
#   // ─── Senda skilaboð til API ────────────────────────────────────
#   // [FASI-2]: /api/chat kallar á MimirCoreV5
#   async function sendMessage(texti) {{
#     try {{
#       const svar = await fetch('/api/chat', {{
#         method: 'POST',
#         headers: {{ 'Content-Type': 'application/json' }},
#         body: JSON.stringify({{
#           message: texti,
#           user_id: 'web_user_fasi1'  // [FASI-3]: Skipta út fyrir JWT user_id
#         }})
#       }});
# 
#       if (!svar.ok) {{
#         const villa = await svar.json().catch(() => ({{}}));
#         throw new Error(villa.detail || `HTTP villa: ${{svar.status}}`);
#       }}
# 
#       const data = await svar.json();
#       return data.response;
# 
#     }} catch (villa) {{
#       console.error('Spjallvilla:', villa);
#       if (villa.message.includes('429')) {{
#         return 'Þú hefur sent of margar beiðnir. Bíddu smástund og reyndu aftur.';
#       }}
#       return `**Villa kom upp:** ${{villa.message}}\\n\\nReyndu aftur eða hafðu samband við okkur.`;
#     }}
#   }}
# 
#   // ─── Aðalfall: meðhöndla sending ──────────────────────────────
#   let isSending = false;
# 
#   async function handleSend() {{
#     if (isSending) return;
# 
#     const inntak = document.getElementById('spjall-inntak');
#     const texti = inntak.value.trim();
#     if (!texti) return;
# 
#     isSending = true;
#     const sendaHnappur = document.getElementById('senda-hnappurinn');
#     sendaHnappur.disabled = true;
#     inntak.disabled = true;
# 
#     // Birta notendaboð
#     const innri = document.getElementById('spjall-innri');
#     const userRow = createUserMsg(texti);
#     innri.appendChild(userRow);
#     inntak.value = '';
#     inntak.style.height = 'auto';
#     scrollBottom();
# 
#     // Birta hlaðningu
#     const {{ row: hlaðningLína, msgText: hlaðningTexti }} = createAiMsg('', true);
#     innri.appendChild(hlaðningLína);
#     scrollBottom();
# 
#     try {{
#       // Sækja svar frá API
#       const svarTexti = await sendMessage(texti);
# 
#       // Skipta hlaðningu út fyrir raunverulegt svar
#       hlaðningTexti.innerHTML = parseMarkdown(svarTexti);
# 
#     }} catch (villa) {{
#       hlaðningTexti.innerHTML = parseMarkdown('**Villa kom upp.** Reyndu aftur.');
#     }} finally {{
#       isSending = false;
#       sendaHnappur.disabled = false;
#       inntak.disabled = false;
#       inntak.focus();
#       scrollBottom();
#     }}
#   }}
# 
#   // ─── Uppsetning þegar DOM er tilbúið ──────────────────────────
#   document.addEventListener('DOMContentLoaded', function() {{
# 
#     // Birta velkomnarboð með lykilsetningum úr mock svörum
#     const velkomin = `Góðan daginn! Ég er **Alvitur**, íslenskur gervigreindaraðstoðarmaður sérhæfður í skjalagreiningu og lögfræðilegri rannsókn.
# 
# Ég get aðstoðað þig með:
# - Greiningu á **ársreikningum** og fjárhagslegum skjölum
# - Leit og samantekt úr **íslenskum lögum og reglugerðum**
# - Samanburð á **samningum** og lagalegum gögnum
# - **Minnisblöð** og skýrslugerð á sekúndum
# 
# Hvað get ég gert fyrir þig í dag?`;
# 
#     const velkoninTexti = document.getElementById('velkomin-texti');
#     const velkoninTimi = document.getElementById('velkomin-timi');
#     if (velkoninTexti) {{
#       velkoninTexti.innerHTML = parseMarkdown(velkomin);
#     }}
#     if (velkoninTimi) {{
#       velkoninTimi.textContent = currentTime();
#     }}
# 
#     // Stillta inntak og takka
#     const inntak = document.getElementById('spjall-inntak');
#     const sendaHnappur = document.getElementById('senda-hnappurinn');
# 
#     // Virkja senda takka þegar texti er til staðar
#     inntak.addEventListener('input', function() {{
#       const hefurTexta = this.value.trim().length > 0;
#       sendaHnappur.disabled = !hefurTexta || isSending;
# 
#       // Sjálfvirk hæðarstilling á textarea
#       this.style.height = 'auto';
#       this.style.height = Math.min(this.scrollHeight, 200) + 'px';
#     }});
# 
#     // Enter sendir, Shift+Enter brýtur row
#     inntak.addEventListener('keydown', function(e) {{
#       if (e.key === 'Enter' && !e.shiftKey) {{
#         e.preventDefault();
#         if (!sendaHnappur.disabled) {{
#           handleSend();
#         }}
#       }}
#     }});
# 
#     // Smella á senda takka
#     sendaHnappur.addEventListener('click', handleSend);
# 
#     // Form submit (mobile vafrar)
#     const chatForm = document.getElementById('chat-form');
#     if (chatForm) {{
#       chatForm.addEventListener('submit', function(e) {{
#         e.preventDefault();
#         if (!sendaHnappur.disabled) {{
#           handleSend();
#         }}
#       }});
#     }}
# 
#     // Touch event á senda takka (mobile)
#     sendaHnappur.addEventListener('touchend', function(e) {{
#       e.preventDefault();
#       if (!this.disabled) {{
#         handleSend();
#       }}
#     }});
# 
#     // Setja fókus á inntak
#     inntak.focus();
#     scrollBottom();
#   }});
#   </script>
# 
# </body>
# </html>"""
# 
# 
# ─── Mock Checkout ────────────────────────────────────────────────────────────

def build_checkout_page(plan: str, amount: int, user_id: str) -> str:
    plan_names = {
        "brons": "Brons",
        "silfur": "Silfur",
        "gull": "Gull",
        "platina": "Platína",
    }
    plan_display = plan_names.get(plan.lower(), plan.capitalize())
    amount_display = f"{amount:,}".replace(",", ".") + " kr" if amount > 0 else "Frítt"

    return """<!DOCTYPE html>
<html lang="is">
<head>
  <title>Greiðsla — {plan_display} — Alvitur</title>
  {SHARED_HEAD}
  {SHARED_STYLES}
</head>
<body>
  <div class="subpage-center">
    <div class="subpage-card">
      <div style="font-size: 2rem; margin-bottom: 0.5rem;">💳</div>
      <h1>Greiðsla</h1>
      <p>Þetta er prófunargreiðslusíða (mock checkout) — engin raunveruleg greiðsla fer fram.</p>

      <div class="checkout-summary">
        <div class="checkout-row">
          <span class="label">Áskrift</span>
          <span class="value">{plan_display}</span>
        </div>
        <div class="checkout-row">
          <span class="label">Notandi</span>
          <span class="value" style="font-family: var(--font-mono); font-size: 0.8125rem;">{user_id}</span>
        </div>
        <div class="checkout-row checkout-total">
          <span class="label">Samtals</span>
          <span class="value">{amount_display} / mán.</span>
        </div>
      </div>

      <form action="/api/webhook/mock_success" method="POST">
        <input type="hidden" name="plan" value="{plan}">
        <input type="hidden" name="amount" value="{amount}">
        <input type="hidden" name="user_id" value="{user_id}">
        <input class="input-field" type="text" placeholder="Nafn á korti" autocomplete="cc-name" required>
        <input class="input-field" type="text" placeholder="0000 0000 0000 0000" autocomplete="cc-number" required>
        <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 0.75rem;">
          <input class="input-field" type="text" placeholder="MM/ÁÁ" autocomplete="cc-exp" required>
          <input class="input-field" type="text" placeholder="CVC" autocomplete="cc-csc" required>
        </div>
        <button type="submit" class="btn btn-primary btn-lg" style="width: 100%; justify-content: center; margin-top: 0.75rem;">
          Staðfesta greiðslu (próf)
        </button>
      </form>

      <a href="/" class="btn btn-secondary" style="width: 100%; justify-content: center; margin-top: 0.75rem;">← Hætta við</a>
    </div>
  </div>
</body>
</html>"""


# ─── Mock Success ─────────────────────────────────────────────────────────────

def build_success_page(plan: str, user_id: str) -> str:
    plan_names = {
        "brons": "Brons",
        "silfur": "Silfur",
        "gull": "Gull",
        "platina": "Platína",
    }
    plan_display = plan_names.get(plan.lower(), plan.capitalize())

    return """<!DOCTYPE html>
<html lang="is">
<head>
  <title>Greiðsla móttekin — Alvitur</title>
  {SHARED_HEAD}
  {SHARED_STYLES}
</head>
<body>
  <div class="subpage-center">
    <div class="subpage-card">
      <div class="success-check">
        <svg width="28" height="28" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round"><path d="M20 6L9 17l-5-5"/></svg>
      </div>
      <h1>Greiðsla móttekin!</h1>
      <p>Þetta er prófunarstaðfesting. Í raun myndi kerfið virkja <strong>{plan_display}</strong> áskrift fyrir notanda <code style="font-family: var(--font-mono); font-size: 0.8125rem; background: var(--bg-elevated); padding: 0.125rem 0.375rem; border-radius: var(--radius-sm);">{user_id}</code>.</p>

      <div class="checkout-summary" style="text-align: left;">
        <div class="checkout-row">
          <span class="label">Staða</span>
          <span class="value" style="color: var(--green);">✓ Staðfest (mock)</span>
        </div>
        <div class="checkout-row">
          <span class="label">Áskrift</span>
          <span class="value">{plan_display}</span>
        </div>
        <div class="checkout-row">
          <span class="label">Notandi</span>
          <span class="value" style="font-family: var(--font-mono); font-size: 0.8125rem;">{user_id}</span>
        </div>
      </div>

      <!-- Sprint 18: /minarsidur skipt út fyrir Telegram -->
      <a href="#" onclick="var el=document.getElementById(\'v5-txt\');if(el)el.scrollIntoView({{behavior:\'smooth\'}});return false;"  target="_blank" rel="noopener noreferrer" class="btn btn-primary btn-lg" style="width: 100%; justify-content: center; margin-bottom: 0.75rem;">
        
      </a>
      <a href="/" class="btn btn-secondary" style="width: 100%; justify-content: center;">← Til baka á forsíðu</a>
    </div>
  </div>
</body>
</html>"""


# ─── Routes ───────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def home():
    """Forsíða — þjónar index.html úr disk."""
    import os
    filepath = os.path.join(os.path.dirname(os.path.abspath(__file__)), "index.html")
    if os.path.exists(filepath):
        with open(filepath, "r", encoding="utf-8") as f:
            return HTMLResponse(content=f.read())
    return HTMLResponse(content=HTML_PAGE)


def build_subpage(titill: str, texti: str) -> str:
    """Sprint 21: Hjálparfall — smíðar einföldu undirsíðu með sama grunn og forsíðan."""
    return HTML_PAGE.replace(
        '<main id="main-content" tabindex="-1">',
        f'<main id="main-content" tabindex="-1"><div class="subpage-card"><h1>{titill}</h1>{texti}</div>',
    ).replace(
        '</main>',
        '</main>',
        1
    )


# ─── Sprint 21: Persónuvernd og Skilmálar ─────────────────────────────

@app.get("/personuvernd", response_class=HTMLResponse)
async def personuvernd():
    """Persónuverndarstefna — Sprint 21, lagatexti frá Aðal & Sigvalda."""
    page = build_subpage(
        titill="Persónuverndarstefna Alviturs",
        texti="""
        <p class="legal-date">Síðast uppfært: 6. apríl 2026</p>
        <p>Hjá Alvitri byggjum við allar okkar lausnir á einni meginreglu: <strong>Þín gögn eru þín, og aðeins þín.</strong> Við skiljum að upplýsingarðu sem þú treystir okkur fyrir, sem greiningarvél fyrir stjórnsýslu og atvinnulífið, eru oft afar viðkvæmar.</p>

        <h2>1. Engin geymsla gagna (Zero-Retention Policy)</h2>
        <p>Þegar þú nýtir skjalagreiningarvél Alviturs (Alvitur Analyze) til að vinna úr skjölum (svo sem PDF, Word eða texta), eru gögnin aðeins unnin í dulkóðuðu vinnsluminni á meðan greining stendur yfir. Um leið og niðurstaða hefur verið afhent, er frumgögnunum tafarlaust eytt af netþjónum okkar. Við geymum aldrei skjöl þín nema beinlínis sé óskað eftir því í lokuðum Enterprise-samningum um varðveislu gagna.</p>

        <h2>2. Engin þjálfun á þínum gögnum (No-Training Policy)</h2>
        <p>Við notum gögn viðskiptavina okkar <strong>ALDREI</strong> til að þjálfa okkar eigin gervigreindarlíkön, né leyfum við þriðja aðila að gera slíkt. Þegar Alvitur nýtir stór erlend API líkön í gegnum bakenda okkar, tryggja “Zero-Data” API-samningar við þá þjónustuaðila að engin inntaksgögn eru notuð í neins konar þjálfunarskyni.</p>

        <h2>3. Gagnaflutningur og dulkóðun</h2>
        <p>Öll samskipti milli þín og netþjóna okkar fara fram bak við lokaða “Zero-Trust” innviði (Cloudflare Tunnels) með fullkominni TLS 1.3 dulkóðun. Engin kerfi okkar eru opin beint út á internetið.</p>

        <p>Fyrir fyrirspurnir varðandi meðferð gagna, eða til að nýta rétt þinn til að gleymast samkvæmt persónuverndarlögum (GDPR), vinsamlegast hafið samband á <a href="mailto:info@alvitur.is">info@alvitur.is</a>.</p>
        """
    )
    return HTMLResponse(content=page)


@app.get("/skilmalar", response_class=HTMLResponse)
async def skilmalar():
    """Notkunarskilmálar — Sprint 21, lagatexti frá Aðal & Sigvalda."""
    page = build_subpage(
        titill="Notkunarskilmálar Alvitur.is",
        texti="""
        <p class="legal-date">Síðast uppfært: 6. apríl 2026</p>
        <p>Með þ ví að heimsækja eða nota þ jónustu Alvitur.is samþykkir þú eftirfarandi skilmála. Alvitur er ætlað sem B2B og B2G greiningartól, hannað til vinnslu gagna, skjala og texta.</p>

        <h2>1. Takmörkun ábyrgðar</h2>
        <p>Alvitur beitir háþróuðum gervigreindarlíkönum til að greina gögn og draga saman samhengi. Notandi samþykkir og gerir sér grein fyrir því að gervigreind er hugbúnaður sem byggir á líkindareikningi og getur sætt svokölluðum “ofskynjunum” (hallucinations) eða rangtúlkunum.</p>
        <p>Allar niðurstöður frá Alvitri eru <strong>EINGÖNGU</strong> til stuðnings við ákvarðanaatöku. Alvitur getur aldrei komið í stað formlegrar faglegar, lögfræðilegrar, læknisfræðilegrar eða fjárhagslegrar ráðgjafar. Notandi ber einn og alfarið alla ábyrgð á því að yfirfara og staðfesta upplýsingar áður en þær eru nýttar. Alvitur.is ehf. ber enga fjárhagslega eða lagalega ábyrgð á beinni eða óbeinni tjóni sem kann að hljótast af notkun þjónustunnar.</p>

        <h2>2. Uppruni gagna og höfundarréttur</h2>
        <p>Notandi ábyrgist að hann hafi fullan og lögmætan rétt til að hlaða upp og vinna með þau gögn sem hann setur í kerfi Alviturs.</p>

        <h2>3. Umgengni og misnotkun</h2>
        <p>Það er með öllu óheimilt að reyna að afhjúpa eða bakfæra (reverse engineer) kerfisarkitektúr Alviturs. Óheimilt er að beita árásum eða inngripum (svo sem “Prompt Injection”) á gervigreindina eða nýta API lykla okkar utan viðskiptasamninga. Slík brot varða umsvifalausri lokun reiknings og hugsanlegum skaðabótakröfum.</p>

        <p>Spurningar: <a href="mailto:info@alvitur.is">info@alvitur.is</a></p>
        """
    )
    return HTMLResponse(content=page)


# @app.get("/minarsidur", response_class=HTMLResponse)
# async def minarsidur():
#     """A) Web Chat Fasi 1 — v1 minimal (sér .js skrá, engin inline JS)"""
#     html_path = "/workspace/mimir_net/static/minarsidur_v1.html"
#     try:
#         with open(html_path, "r", encoding="utf-8") as f:
#             return HTMLResponse(content=f.read())
#     except FileNotFoundError:
#         return HTMLResponse(content=build_minarsidur_page())
# 


@app.get("/api/health")
async def health():
    """Heilsufarsskoðun — notað af monitoring og load balancer."""
    return JSONResponse(content={
        "status": "ok",
        "version": "sprint63-track-b",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "fasi": "production",
    })

@app.get("/api/health/detailed")
async def health_detailed():
    """Raunverulegur health check fyrir alla Alvitur components."""
    import time, subprocess, os
    import httpx
    from datetime import datetime, timedelta
    start = time.time()

    report = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "overall": "unknown",
        "components": {},
        "metrics": {},
        "errors_last_10min": 0,
    }

    # 1. FastAPI self
    report["components"]["fastapi"] = {
        "status": "ok",
        "pid": os.getpid(),
        "version": "sprint63-track-b",
    }

    # 2. vLLM local (port 8002)
    try:
        async with httpx.AsyncClient(timeout=3.0) as client:
            r = await client.get("http://localhost:8002/v1/models")
            if r.status_code == 200:
                models = r.json().get("data", [])
                report["components"]["vllm"] = {
                    "status": "ok",
                    "models_loaded": [m.get("id") for m in models[:3]],
                    "latency_ms": int((time.time() - start) * 1000),
                }
            else:
                report["components"]["vllm"] = {"status": "degraded", "http_code": r.status_code}
    except Exception as e:
        report["components"]["vllm"] = {"status": "down", "error": f"{type(e).__name__}: {e}"}

    # 3. OpenRouter
    try:
        key = os.environ.get("OPENROUTER_API_KEY", "")
        if not key or len(key) < 40:
            report["components"]["openrouter"] = {"status": "misconfigured", "error": "API key missing"}
        else:
            async with httpx.AsyncClient(timeout=5.0) as client:
                r = await client.get(
                    "https://openrouter.ai/api/v1/auth/key",
                    headers={"Authorization": f"Bearer {key}"},
                )
                if r.status_code == 200:
                    data = r.json().get("data", {})
                    report["components"]["openrouter"] = {
                        "status": "ok",
                        "credits_remaining": data.get("limit_remaining"),
                        "rate_limit": data.get("rate_limit", {}).get("requests"),
                    }
                else:
                    report["components"]["openrouter"] = {"status": "auth_failed", "http_code": r.status_code}
    except Exception as e:
        report["components"]["openrouter"] = {"status": "unreachable", "error": f"{type(e).__name__}: {e}"}

    # 4. Disk space
    try:
        result = subprocess.run(["df", "-h", "/workspace"], capture_output=True, text=True, timeout=5)
        lines = result.stdout.strip().split("\n")
        if len(lines) >= 2:
            parts = lines[1].split()
            pct = int(parts[4].rstrip("%"))
            report["components"]["disk"] = {
                "status": "ok" if pct < 90 else "warning",
                "used": parts[2],
                "available": parts[3],
                "percent_used": parts[4],
            }
    except Exception as e:
        report["components"]["disk"] = {"status": "error", "error": str(e)}

    # 5. GPU VRAM
    try:
        result = subprocess.run(
            ["nvidia-smi", "--query-gpu=memory.used,memory.total,memory.free", "--format=csv,noheader,nounits"],
            capture_output=True, text=True, timeout=3
        )
        parts = result.stdout.strip().split(", ")
        if len(parts) == 3:
            used, total, free = int(parts[0]), int(parts[1]), int(parts[2])
            pct = (used / total) * 100
            report["components"]["gpu"] = {
                "status": "ok" if pct < 95 else "critical",
                "used_mib": used,
                "total_mib": total,
                "free_mib": free,
                "percent_used": round(pct, 1),
            }
    except Exception as e:
        report["components"]["gpu"] = {"status": "error", "error": str(e)}

    # 6. Error count sidastu 10 min
    try:
        result = subprocess.run(
            ["tail", "-n", "500", "/workspace/web_server.log"],
            capture_output=True, text=True, timeout=3
        )
        error_count = sum(
            1 for line in result.stdout.split("\n")
            if any(k in line for k in ("ERROR", "Traceback", "Exception"))
        )
        report["errors_last_10min"] = error_count
    except Exception:
        report["errors_last_10min"] = -1

    # 7. Overall status
    statuses = [c.get("status", "unknown") for c in report["components"].values()]
    if any(s == "down" for s in statuses):
        report["overall"] = "critical"
    elif any(s in ("degraded", "warning", "error", "critical") for s in statuses):
        report["overall"] = "degraded"
    elif all(s == "ok" for s in statuses):
        report["overall"] = "healthy"
    else:
        report["overall"] = "unknown"

    report["metrics"]["check_latency_ms"] = int((time.time() - start) * 1000)
    return report


@app.get("/api/diagnostics")
async def diagnostics():
    """Sprint 63 Track A5 + B3: Diagnostics — stöðu leiða og umhverfis."""
    import os as _os_d
    import time as _time_d
    _key = _os_d.environ.get("OPENROUTER_API_KEY", "")
    leid_a_enabled = bool(_key and len(_key) > 10 and not _key.startswith("sk-or-v1-BAD"))
    # Track B3: vLLM er á port 8002 (ekki 8001 sem var hardcoded default)
    _sovereign_url = _os_d.environ.get("SOVEREIGN_URL", "http://localhost:8002/v1/chat/completions")
    # Track B3: env flag + port
    _env_flag = _os_d.environ.get("ALVITUR_ENV", "prod")
    _port = int(_os_d.environ.get("ALVITUR_PORT", "8000"))
    try:
        uptime = int(_time_d.time() - _SERVER_START_TIME)
    except NameError:
        uptime = -1
    return JSONResponse(content={
        "status": "ok",
        "version": "sprint63-track-b",
        "env": _env_flag,
        "port": _port,
        "leid_a_enabled": leid_a_enabled,
        "leid_a_key_length": len(_key) if _key else 0,
        "leid_b_enabled": bool(_sovereign_url),
        "leid_b_url": _sovereign_url,
        "uptime_seconds": uptime,
        "loaded_env": bool(_key),
        "timestamp": datetime.now(timezone.utc).isoformat(),
    })


# ─── Sprint 58+59: MCP Tools Endpoints ─────────────────────────────────────────

@app.get("/api/tools")
async def tools_list():
    """
    Sprint 59: Skilar lista af öllum tiltækum tools.
    MCP-samhæft — notað af MCP clients og /api/tools wiring.
    """
    from interfaces.mcp_server import mcp_list_tools
    tools = await mcp_list_tools()
    return JSONResponse(content={"success": True, "tools": tools, "count": len(tools)})


@app.post("/api/tools/{tool_name}")
async def tools_call(tool_name: str, request: Request):
    """
    Sprint 59: Kallar á tool með gefnum arguments.
    Body: JSON með arguments fyrir tool.
    Skilar niðurstöðu frá tool.
    """
    from interfaces.mcp_server import mcp_call_tool
    try:
        body = await request.json()
    except Exception:
        body = {}
    result = await mcp_call_tool(tool_name, body)
    status = 200 if result.get("success") else 404 if "ekki til" in result.get("error", "") else 502
    return JSONResponse(content=result, status_code=status)


# ─── Sprint 21: PDF Analyze Endpoint ──────────────────────────────────────────────

# ── Sprint 28: K1/K2 — .docx and .xlsx parsers ───────────────────────────────
MAX_DOC_SIZE = 20 * 1024 * 1024  # 20 MB (same as PDF)

MAGIC_BYTES = {
    b'%PDF': 'pdf',
    b'PK\x03\x04': 'office',  # .docx and .xlsx are ZIP-based Office formats
}

def _detect_filetype(data: bytes, filename: str) -> str:
    """Return 'pdf', 'docx', 'xlsx', or raise HTTPException."""
    ext = filename.lower().rsplit('.', 1)[-1] if '.' in filename else ''
    header = data[:4]
    if header == b'%PDF':
        if ext != 'pdf':
            raise HTTPException(status_code=415,
                detail="Skráin er merkt sem PDF en innihald stemmir ekki.")
        return 'pdf'
    if header[:2] == b'PK':
        if ext == 'docx':
            return 'docx'
        if ext == 'xlsx':
            return 'xlsx'
        raise HTTPException(status_code=415,
            detail="Office skjal þekkt en skráarending er óþêkkt. Sendu .docx eða .xlsx.")
    raise HTTPException(status_code=415,
        detail="Skráargerð ekki stuðd. Styður PDF, Word (.docx) og Excel (.xlsx).")


def _parse_docx(data: bytes) -> tuple[int, list[str]]:
    """Extract text from .docx. Returns (page_estimate, text_parts)."""
    from docx import Document
    import io
    doc = Document(io.BytesIO(data))
    parts = []
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t:
            parts.append(t)
    # Estimate pages: ~3000 chars per page
    total_chars = sum(len(p) for p in parts)
    page_estimate = max(1, total_chars // 3000)
    return page_estimate, parts


def _parse_xlsx(data: bytes) -> tuple[int, list[str]]:
    """Extract text from .xlsx. Returns (sheet_count, text_parts)."""
    import openpyxl, io
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    sheet_count = len(wb.sheetnames)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text = []
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                rows_text.append(" | ".join(cells))
                row_count += 1
            if row_count >= 500:  # cap at 500 rows per sheet
                rows_text.append("[... fleiri raðir ...]")
                break
        if rows_text:
            parts.append(f"[Blað: {sheet_name}]\n" + "\n".join(rows_text))
    wb.close()
    return sheet_count, parts

# ─────────────────────────────────────────────────────────────────────────────
# ─ Sprint 27: S4 Wallet Circuit Breaker ───────────────────────────────────
WALLET_MIN_USD       = float(os.environ.get("WALLET_MIN_USD", "5.0"))
WALLET_MIN_VAULT_USD = float(os.environ.get("WALLET_MIN_VAULT_USD", "10.0"))

# PLG quota tracker (IP-based, in-memory)
_quota_tracker_chat: dict = {}  # /api/chat quota per IP
_quota_tracker_doc: dict = {}   # /api/analyze-document quota per IP
FREE_QUOTA = 5



# -- Sprint 66 pre-A hotfix: persist _beta_tracker across restarts --
import json as _json_bt
import os as _os_bt
import time as _time_bt
from pathlib import Path as _Path_bt

_BETA_TRACKER_FILE = _Path_bt(_os_bt.getenv("BETA_TRACKER_PATH", "data/beta_tracker.json"))

def _load_beta_tracker_from_disk() -> dict:
    try:
        if not _BETA_TRACKER_FILE.exists():
            return {}
        raw = _json_bt.loads(_BETA_TRACKER_FILE.read_text(encoding="utf-8"))
        now = _time_bt.time()
        # TODO(S66-A): replace with BETA_DURATION_SEC module const (defined below)
        _DUR = int(_os_bt.getenv("BETA_DURATION_SEC_OVERRIDE", 7 * 24 * 3600))
        return {ip: float(ts) for ip, ts in raw.items() if now - float(ts) <= _DUR}
    except Exception as e:
        try:
            logger.warning(f"[BETA] load failed: {type(e).__name__}: {e}")
        except Exception:
            pass
        return {}

def _save_beta_tracker_to_disk(tracker: dict) -> None:
    try:
        _BETA_TRACKER_FILE.parent.mkdir(parents=True, exist_ok=True)
        tmp = _BETA_TRACKER_FILE.with_suffix(".tmp")
        tmp.write_text(_json_bt.dumps(tracker, indent=2), encoding="utf-8")
        tmp.replace(_BETA_TRACKER_FILE)
    except Exception as e:
        try:
            logger.warning(f"[BETA] save failed: {type(e).__name__}: {e}")
        except Exception:
            pass
# -- /persist helpers --

# ── Sprint 62: Beta tracker ──
# Beta-tier via human phrase in chat: "Sigvaldi sendi mig" (7-day renewable)
_beta_tracker: dict[str, float] = {}   # IP → promotion_timestamp (epoch sec)
BETA_DURATION_SEC = 7 * 24 * 3600      # 7 dagar
BETA_FRASAR = (
    "sigvaldi sendi mig",
    "ég er beta notandi",
    "beta aðgangur",
    "beta adgangur",
)

def _er_beta_fras(text: str) -> bool:
    """Satt ef notendatexti inniheldur einhvern viðurkenndan beta-frasa."""
    if not text:
        return False
    lower = text.lower()
    return any(fras in lower for fras in BETA_FRASAR)

def _er_beta_ip(ip: str) -> bool:
    """Satt ef IP er í beta-tier (innan 7 daga frá promotion)."""
    import time as _t
    ts = _beta_tracker.get(ip)
    if ts is None:
        return False
    if _t.time() - ts > BETA_DURATION_SEC:
        _beta_tracker.pop(ip, None)
        return False
    return True

def _promota_beta(ip: str) -> None:
    """Setur IP í beta-tier núna (eða endurnýjar)."""
    import time as _t
    _beta_tracker[ip] = _t.time()
    _save_beta_tracker_to_disk(_beta_tracker)
# ── /Sprint 62 Beta tracker ──
_beta_tracker.update(_load_beta_tracker_from_disk())
_wallet_cache: dict  = {"balance": None, "ts": 0.0}
_WALLET_TTL          = 120

def _get_openrouter_balance():
    import time as _t, requests as _rq
    now = _t.time()
    if _wallet_cache["balance"] is not None and now - _wallet_cache["ts"] < _WALLET_TTL:
        return _wallet_cache["balance"]
    try:
        _k = os.environ.get("OPENROUTER_API_KEY", "")
        r = _rq.get("https://openrouter.ai/api/v1/auth/key",
                    headers={"Authorization": f"Bearer {_k}"}, timeout=5)
        if r.status_code == 200:
            d = r.json().get("data", {})
            limit = d.get("limit")
            if limit is not None:
                bal = round(float(limit) - float(d.get("usage") or 0), 4)
            else:
                bal = 100.0
            _wallet_cache["balance"] = bal
            _wallet_cache["ts"] = now
            return bal
    except Exception:
        pass
    return None

def _wallet_preflight(is_vault=False):
    bal = _get_openrouter_balance()
    if bal is None: return
    threshold = WALLET_MIN_VAULT_USD if is_vault else WALLET_MIN_USD
    if bal < threshold:
        logger.warning(f"Wallet circuit breaker: balance=${bal:.2f} < ${threshold:.2f}")
        raise HTTPException(status_code=503,
            detail="Þ jónustan er tímabundið ótilðæk. Reyndu aftur eftir augnablik.")

# ───────────────────────────────────────────────────────────
SECURE_DOCS_DIR = Path("/workspace/mimir_net/secure_docs")
SECURE_DOCS_DIR.mkdir(parents=True, exist_ok=True)
MAX_PDF_SIZE = 20 * 1024 * 1024  # Sprint 27 S2: raised to 20 MB



def _log_intent(endpoint: str, query, filename, file_size, tier) -> None:
    """Sprint 64 B2-V2: observability hook. NEVER raises."""
    if not _INTENT_AVAILABLE or _classify_intent is None:
        return
    try:
        ir = _classify_intent(query=query, filename=filename,
                              file_size=file_size, tier=tier)
        logger.info(
            f"[INTENT] endpoint={endpoint} domain={ir.domain} "
            f"depth={ir.reasoning_depth} conf={ir.confidence_score:.2f} "
            f"sens={ir.sensitivity} adapter={ir.adapter_hint} "
            f"src={ir.source_hint}"
        )
    except Exception as _ie:
        logger.warning(f"[INTENT] classify failed on {endpoint}: {type(_ie).__name__}: {_ie}")

@app.post("/api/analyze-document")
async def analyze_document(request: Request, file: Optional[UploadFile] = File(None), query: Optional[str] = Form(None)):
    """
    Sprint 21 — PDF Analyze Endpoint (Lag 1).
    Tekur við PDF skrá, les texta með PyMuPDF og búr til grófótta greiningu.
    Zero-Data: Skrá er eyð um leið og vinnslu ljúkur.
    Sprint 43b: query parameter wired from multipart form data.
    Sprint 43c: file optional — text-only queries fall back to direct LLM call.
    """

    # ── Sprint 43c: text-only fallback ──────────────────────────────
    _tier_hdr = request.headers.get("X-Alvitur-Tier", "general").lower().strip() if request else "general"
    _llm_model = _MODEL_LEIDA_B if _tier_hdr == "vault" else _MODEL_LEIDA_A
    if file is None or file.filename == "":
        if not query or not query.strip():
            return JSONResponse(status_code=422, content={
                "error_code": "empty_prompt",
                "detail": "Vinsamlegast skrifaðu fyrirspurn eða hladdu upp skjali.",
            })
        # Sprint 46 Phase 1b: quota check á text-only path (var vantar)
        import hashlib as _hl
        _master_key_txt = os.environ.get("ALVITUR_MASTER_KEY_HASH", "")
        _req_key_txt = request.headers.get("X-Master-Key", "") if request else ""
        _is_admin_txt = bool(_master_key_txt and _req_key_txt and _hl.sha256(_req_key_txt.encode()).hexdigest() == _master_key_txt)
        _client_ip_txt = (
            request.headers.get("CF-Connecting-IP")
            or request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
            or (request.client.host if request.client else "unknown")
        )
        # ── Sprint 64 A1: Beta check á text-only path (var vantar) ──
        # Speglarar logic í /api/chat (line ~3593) og analyze-document file-path (line ~3294)
        if _er_beta_fras(query or ""):
            _promota_beta(_client_ip_txt)
            logger.info(f"[BETA] {_client_ip_txt} promotaður í beta-tier (7d) via text-only")
        _is_beta_txt = _er_beta_ip(_client_ip_txt)
        # ── /Sprint 64 A1 beta check ──
        _log_intent("analyze-document/text-only", query, None, None, _tier_hdr)
        _quota_count_txt = _quota_tracker_doc.get(_client_ip_txt, 0) + 1
        if not _is_admin_txt and not _is_beta_txt:
            _quota_tracker_doc[_client_ip_txt] = _quota_count_txt
        if _quota_count_txt > FREE_QUOTA and not _is_admin_txt and not _is_beta_txt:
            return JSONResponse(status_code=403, content={
                "success": False,
                "error": "Ókeypis prufutími er liðinn.",
                "error_code": "quota_exceeded",
                "upgrade_required": True,
                "upgrade_url": "/askrift",
                "message": "Þú hefur nýtt þér 5 ókeypis beiðnir. Uppfærðu til að halda áfram.",
            })

        # Text-only path: send query directly to LLM (async — Sprint 46 Phase 1b)
        import os as _os
        import httpx as _httpx
        from datetime import datetime as _dt, timezone as _tz
        _key = _os.environ.get("OPENROUTER_API_KEY", "")
        _tier = _tier_hdr
        _summary = None
        # 🟢 Sprint 62 Patch E v2: no-key → sovereign Leið B strax (skip OpenRouter entirely)
        if not _key and _tier_hdr != "vault":
            logger.warning("[ALVITUR] Sprint62E NO-KEY → sovereign fallback á Leið B")
            try:
                async with _VAULT_SEMAPHORE_WS:
                    _summary, _model_used, _usage = await _call_leid_b((query or "").strip())
            except Exception as _fe:
                logger.error(f"[ALVITUR] Sprint62E exc: {type(_fe).__name__}: {_fe}")
                _summary = None
            if _summary is None:
                return JSONResponse(status_code=503, content={
                    "error_code": "sovereign_unavailable",
                    "detail": "Staðbundin gervigreind er að ræsast. Reyndu eftir andartak."})
            logger.info(f"[ALVITUR] Sprint62E sovereign OK model={_model_used}")
            return JSONResponse(status_code=200, content={
                "success": True, "summary": _summary, "mode": "text-only",
                "pipeline_source": f"sovereign_nokey_{_model_used}", "model": _model_used})
        if _key:
            try:
                _now_str = _dt.now(_tz.utc).strftime("%Y-%m-%d %H:%M UTC")
                # Sprint 47: Domain classification + specialist prompt
                # Sprint 60c: None guard + honesty instruction
                from interfaces.specialist_prompts import get_specialist_prompt as _get_prompt
                from interfaces.skills.classify import ClassifySkill as _ClsSkill
                try:
                    _domain_txt = await _ClsSkill().run((query or "").strip() or "general", tier=_tier_hdr)
                except Exception:
                    _domain_txt = "general"
                _domain_txt = _domain_txt or "general"
                # s68-hotfix Part 3: text-only + tone-guide + thinking-suppress
                _honesty = (
                    "\n\nMikilvægt: Notandi hefur ekki hengt við skjal — þetta er almenn "
                    "spurning. Svaraðu út frá almennri þekkingu þinni á íslensku. "
                    "Ef þú veist ekki svarið með vissu, segðu það heiðarlega og bjóddu "
                    "framhaldsspurningu. Ekki búa til staðreyndir. "
                    "MIKILVÆGT UM FRAMSETNINGU: Ekki sýna hugsanaferli þitt — engar "
                    "<thinking>-tög, engin ‚bíð, ég þarf að“ setningar, engin internal-monologue. "
                    "Sendu aðeins hreint, prófessjónal svar beint. Engar emoji, "
                    "engin upphrópunarmerki, formlegt B2B-mál."
                )
                _system_prompt = _get_prompt(_domain_txt, _now_str) + _honesty
                # Sprint 70 Track D — RAG+ hook (text-only path)
                try:
                    from core.rag_orchestrator import retrieve_legal_context, build_rag_injection
                    _rag_txt = retrieve_legal_context(
                        query=(query or "").strip(),
                        intent_domain=_domain_txt,
                        tier=_tier,
                        tenant_id="system",
                    )
                    if _rag_txt.refusal:
                        from fastapi.responses import JSONResponse as _JR
                        return _JR(content={"success":True,"response":_rag_txt.refusal,
                            "pipeline_source":"rag_refusal_vault","domain":_domain_txt,
                            "zero_data":True,"found":True,"status":"ready_for_analysis",
                            "citations":[],"quota_warning":None,"tier":_tier})
                    if _rag_txt.used_retrieval:
                        _system_prompt = _system_prompt + build_rag_injection(_rag_txt.chunks)
                        _pipeline_source_txt = "rag_grounded_" + _tier
                    elif _rag_txt.fallback_to_gemini:
                        _system_prompt = _system_prompt + "[ATH: Engin lagatilvitnun fannst. Svaraðu varlega.]"
                        _pipeline_source_txt = "rag_fallback_general"
                except Exception as _rag_e2:
                    logger.warning("[RAG] text-only villa: %s", _rag_e2)
                logger.info(f"[ALVITUR] Sprint61 text-only tier={_tier} domain={_domain_txt}")
                _pipeline_source_txt = "unknown"
                if _tier == "vault":
                    from interfaces.config import VAULT_MAX_INPUT_TOKENS as _vmax
                    if _estimate_tokens(query or "") > _vmax:
                        return JSONResponse(status_code=413, content={
                            "error_code": "vault_input_too_large",
                            "detail": "Fyrirspurn er of stór fyrir Vault tier (max 8000 tokens). Styttu textann."})
                    async with _VAULT_SEMAPHORE_WS:
                        _summary, _model_used, _usage = await _call_leid_b(query.strip())
                    if _summary is None:
                        return JSONResponse(status_code=503, content={
                            "error_code": "vault_local_unavailable",
                            "detail": "Trúnaðarþjónusta tímabundið ekki tiltæk. Local AI module er að ræsast — reyndu aftur eftir 1 mínútu."})
                    _pipeline_source_txt = f"local_vllm_{_model_used}"
                    _in_tok = _usage.get("prompt_tokens", 0); _out_tok = _usage.get("completion_tokens", 0)
                    logger.info(f"[ALVITUR] leid_b sovereign tokens in={_in_tok} out={_out_tok}")
                else:
                    _summary, _model_used, _usage = await _call_leid_a(_system_prompt, query.strip())
                    if _summary is None:
                        logger.warning("[ALVITUR] Sprint62C text-only: Leið A None, reyni fallback á Leið B")
                        try:
                            async with _VAULT_SEMAPHORE_WS:
                                _summary, _model_used, _usage = await _call_leid_b(query.strip())
                        except Exception as _fe:
                            logger.error(f"[ALVITUR] Sprint62C fallback exc: {type(_fe).__name__}: {_fe}")
                            _summary = None
                        if _summary is None:
                            return JSONResponse(status_code=503, content={
                                "error_code": "both_pipelines_unavailable",
                                "detail": "Þjónusta tímabundið ekki aðgengileg. Reyndu eftir augnablik."})
                        _pipeline_source_txt = f"fallback_local_{_model_used}"
                        logger.info(f"[ALVITUR] Sprint62C fallback OK: model={_model_used}")
                    else:
                        _pipeline_source_txt = f"openrouter_{_model_used.split('/')[-1]}"
                    _in_tok = _usage.get("prompt_tokens", 0); _out_tok = _usage.get("completion_tokens", 0)
                    _cost = (_in_tok * 0.15 + _out_tok * 0.60) / 1_000_000 if "gpt-4o-mini" in _model_used else (_in_tok * 3.00 + _out_tok * 15.00) / 1_000_000 if "claude-sonnet" in _model_used else 0.0
                    logger.info("[ALVITUR] token_obs pipeline=text_query model=%s tier=%s in=%d out=%d cost=%.6f", _model_used, _tier, _in_tok, _out_tok, _cost)
                    try:
                        # from interfaces.chat_routes import _polish as _polish_fn_txt
                        _summary = await _polish_fn_txt(_summary, _key)
                    except Exception as _pe:
                        logger.warning(f"[ALVITUR] polish failed (non-fatal): {_pe}")
            except Exception as _e:
                logger.error(f"[ALVITUR] text-only pipeline exc: {type(_e).__name__}: {_e}")
                _summary = None
                _domain_txt = "general"
                _pipeline_source_txt = "error"
        if _summary is None:
            # Sprint 62 Patch C: Final safety net — both pipelines down
            return JSONResponse(status_code=503, content={
                "error_code": "service_unavailable",
                "detail": "Greiningar þjónusta tímabundið ekki aðgengileg. Reyndu aftur eftir andartak.",
            })
        return JSONResponse(content={
            "success": True,
            "filename": None,
            "sidur": 0,
            "zero_data": True,
            "tier": _tier,
            "query": query.strip(),
            "response": _summary,
            "domain": _domain_txt,
            "citations": [],
            "found": True,
            "status": "ready_for_analysis",
            "pipeline_source": _pipeline_source_txt if "_pipeline_source_txt" in dir() else "unknown",
            "rag_metadata": _rag_txt.__dict__ if hasattr(_rag_txt,"used_retrieval") else {},
        })

    import fitz  # PyMuPDF

    # — 1. Staðfesta skrártegund —
    _ext = file.filename.lower().rsplit(".", 1)[-1] if "." in file.filename else ""
    if _ext not in ("pdf", "docx", "xlsx"):
        raise HTTPException(status_code=400, detail="Skráargerð ekki stuð. Sendu PDF, Word eða Excel.")

    # S2: Stream size check before full read
    _cl = None
    if hasattr(file, "headers") and file.headers:
        _cl = file.headers.get("content-length")
    if _cl and int(_cl) > MAX_PDF_SIZE:
        raise HTTPException(status_code=413,
            detail="Skjalið er yfir 20 MB — minnkaðu eða veldu annað.")
    efni = await file.read(MAX_PDF_SIZE + 1)
    if len(efni) > MAX_PDF_SIZE:
        raise HTTPException(status_code=413,
            detail="Skjalið er yfir 20 MB — minnkaðu eða veldu annað.")

    # K1/K2/S3: Filetype detection (magic bytes + extension)
    _filetype = _detect_filetype(efni, file.filename)


    # S4: Wallet circuit breaker
    _tier = request.headers.get("X-Alvitur-Tier", "general").lower().strip() if request else "general"
    _is_vault = (_tier == "vault")
    _wallet_preflight(is_vault=_is_vault)

    # PLG quota check — master key bypass
    _master_key = os.environ.get("ALVITUR_MASTER_KEY_HASH", "")
    _req_key = request.headers.get("X-Master-Key", "") if request else ""
    import hashlib as _hl
    _is_admin = bool(_master_key and _req_key and _hl.sha256(_req_key.encode()).hexdigest() == _master_key)
    _client_ip = (request.headers.get("CF-Connecting-IP") or request.headers.get("X-Forwarded-For", "").split(",")[0].strip() or (request.client.host if request.client else "unknown"))
    # Sprint 62 Patch A.1: define _is_beta early (was defined 240 lines later, causing NameError)
    try:
        _is_beta = _er_beta_ip(_client_ip)
    except Exception:
        _is_beta = False
    try:
        _fn = file.filename if file else None
        _fsize = getattr(file, "size", None) if file else None
    except Exception:
        _fn, _fsize = None, None
    _log_intent("analyze-document/file", query, _fn, _fsize,
                request.headers.get("X-Alvitur-Tier", "general") if request else "general")
    _quota_count = _quota_tracker_doc.get(_client_ip, 0) + 1
    if not _is_admin:
        _quota_tracker_doc[_client_ip] = _quota_count
    if _quota_count > FREE_QUOTA and not _is_admin and not _is_beta:
        return JSONResponse(status_code=403, content={
            "success": False,
            "error": "Ókeypis prufutími er liðinn.",
            "error_code": "quota_exceeded",
            "upgrade_required": True,
            "upgrade_url": "/askrift",
            "message": "Þú hefur nýtt þér 5 ókeypis beiðnir. Uppfærðu til að halda áfram.",
        })

    # PLG warning
    _quota_warning = None
    if _quota_count == 4 and not _is_admin:
        _quota_warning = "Þú hefur 1 beiðni eftir af 5 ókeypis beiðnum."
    elif _quota_count == 5 and not _is_admin:
        _quota_warning = "Þetta er þín síðasta ókeypis beiðni. Uppfærðu til að halda áfram."

    # S5+S6: Unified Parser (Office + PDF)
    # Office files (XLSX/DOCX)
    if _filetype in ('xlsx', 'docx'):
        import io
        try:
            if _filetype == 'xlsx':
                # Sprint 62 Patch B: Pandas pre-processor (Reiknivélar-Agent).
                # Returns markdown with real computed sums so LLM doesn't hallucinate.
                try:
                    from interfaces.excel_preprocessor import preprocess_excel as _prep_xlsx
                    heildartexti = _prep_xlsx(efni)
                    sidur = 1
                    logger.info(f"[ALVITUR] Sprint62b xlsx preprocessed via pandas, {len(heildartexti)} chars")
                except Exception as _xe:
                    logger.warning(f"[ALVITUR] pandas preprocess failed, fallback to openpyxl flat: {type(_xe).__name__}: {_xe}")
                    import openpyxl
                    wb = openpyxl.load_workbook(io.BytesIO(efni), read_only=True, data_only=True)
                    txt = []
                    for ws in wb.worksheets:
                        for row in ws.iter_rows(values_only=True):
                            clean = [str(c) if c is not None else "" for c in row]
                            line = " | ".join(v for v in clean if v.strip())
                            if line: txt.append(line)
                    heildartexti = "\n".join(txt)
                    sidur = len(wb.worksheets)
            elif _filetype == 'docx':
                from docx import Document
                doc = Document(io.BytesIO(efni))
                heildartexti = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                sidur = len(doc.paragraphs)
        except Exception as e:
            heildartexti = f"Villa við lestur: {str(e)}"
            sidur = 0
    else:
        # PDF Logic (Original)
        import concurrent.futures as _cf
        
        def _parse_pdf(data):
            parts, pages = [], 0
            with fitz.open(stream=data, filetype="pdf") as doc:
                pages = len(doc)
                for idx, pg in enumerate(doc):
                    t = pg.get_text().strip()
                    if t: parts.append(f"[Síða {idx+1}]\n{t}")
            return pages, parts

        texti_hlutar, sidur = [], 0
        if _is_vault:
            # VAULT: zero disk write
            with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                fut = ex.submit(_parse_pdf, efni)
                try:
                    sidur, texti_hlutar = fut.result(timeout=90)
                except _cf.TimeoutError:
                    raise HTTPException(status_code=504, detail="Vinnsla tók of langan tíma.")
        else:
            # GENERAL: disk write + cleanup
            skra_id = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S_%f")
            skra_slod = SECURE_DOCS_DIR / f"doc_{skra_id}.pdf"
            try:
                skra_slod.write_bytes(efni)
                with _cf.ThreadPoolExecutor(max_workers=1) as ex:
                    fut = ex.submit(_parse_pdf, efni)
                    try:
                        sidur, texti_hlutar = fut.result(timeout=90)
                    except _cf.TimeoutError:
                        raise HTTPException(status_code=504, detail="Vinnsla tók of langan tíma.")
            finally:
                if skra_slod.exists(): skra_slod.unlink()
        
        heildartexti = "\n\n".join(texti_hlutar)

    if not heildartexti.strip():
        return JSONResponse(content={
            "success": False,
            "error": "Ekki tókst að lesa texta úr PDF (gæti verið skannað mynd).",
            "sidur": sidur,
        })

    # — 6. Skila grunnupplýsingum (LLM greining kemur í Lag 2) —
    # Sprint 62 Patch A: Only run fitz (PDF parser) for PDF files.
    # xlsx/docx already parsed above via openpyxl/python-docx.
    _parts = []
    if _filetype == 'pdf':
        import fitz as _fitz_inner
        with _fitz_inner.open(stream=efni, filetype="pdf") as _doc:
            sidur = len(_doc)
            for _pg in _doc:
                _t = _pg.get_text().strip()
                if _t: _parts.append(_t)
    # Default result fallback for all filetypes
    result = {"response": heildartexti if "heildartexti" in dir() else "", "citations": [], "found": bool(heildartexti.strip()) if "heildartexti" in dir() else False}

    # Build result dict for PDF fallback
    if _filetype == 'pdf':
        result = {"response": " ".join(_parts[:3000]), "citations": [], "found": bool(_parts)}
    import os as _os
    _tier = request.headers.get("X-Alvitur-Tier", "general").lower()
    _summary = None
    _pipeline_source_doc = "unknown"
    _domain_doc = "general"
    _rag_doc_meta = {"used_retrieval": False, "chunks_count": 0, "top_score": 0.0, "low_confidence": False, "source_laws": []}
    if _tier in ("general", "vault"):
        try:
            # 🟢 Sprint 62 Patch F: no-key → sovereign Leið B fyrir file-upload
            _key_preflight = _os.environ.get("OPENROUTER_API_KEY", "")
            if not _key_preflight and _tier != "vault":
                logger.warning("[ALVITUR] Sprint62F file-upload NO-KEY → sovereign Leið B")
                _msg_fallback = f"SPURNING: {(query or '').strip() or 'Greindu þetta skjal.'}\n\nSKJAL:\n{heildartexti[:30000]}"
                try:
                    async with _VAULT_SEMAPHORE_WS:
                        _summary, _model_used, _usage = await _call_leid_b(_msg_fallback)
                except Exception as _fe:
                    logger.error(f"[ALVITUR] Sprint62F exc: {type(_fe).__name__}: {_fe}")
                    _summary = None
                if _summary is not None:
                    _pipeline_source_doc = f"sovereign_nokey_file_{_model_used}"
                    logger.info(f"[ALVITUR] Sprint62F sovereign OK model={_model_used}")
                    # 🟢 Sprint 62 Patch I: EARLY RETURN — ekki keyra leið A aftur!
                    _in_tok = _usage.get("prompt_tokens", 0); _out_tok = _usage.get("completion_tokens", 0)
                    logger.info(f"[ALVITUR] Sprint62I early-return (skip leid_a) in={_in_tok} out={_out_tok}")
                    return JSONResponse(status_code=200, content={
                        "success": True,
                        "filename": _filename if "_filename" in dir() else "",
                        "sidur": sidur if "sidur" in dir() else 1,
                        "zero_data": True,
                        "tier": _tier,
                        "query": (query or "").strip(),
                        "response": _summary,
                        "domain": "general",
                        "citations": [],
                        "found": True,
                        "status": "ready_for_analysis",
                        "quota_warning": None,
                        "pipeline_source": _pipeline_source_doc,
                    })
            _key = _os.environ.get("OPENROUTER_API_KEY", "")
            _context_limit = 60000 if _tier == "vault" else 30000
            _full = heildartexti[:_context_limit]
            if _full:
                if query and query.strip():
                    truncation_note = " [Skjal styttur vegna stærðar]" if len(heildartexti) > _context_limit else ""
                    _msg = f"""SPURNING: {query.strip()}

REGLUR:
1. Svaraðu BEINT spurningunni hér að ofan á íslensku.
2. Notaðu skjalið AÐEINS sem heimild til að styðja við svarið.
3. EKKI endurskrifa skjalið, EKKI draga það saman nema beðið sé um það sérstaklega.
4. Ef upplýsingar vantar í skjalinu, segðu það hreint út.

SKJAL{truncation_note}:
{_full}"""
                else:
                    _msg = f"""Greindu þetta skjal stuttlega á íslensku (max 3 setningar).
Áhersla á meginþætti og ályktanir.

SKJAL:
{_full[:3000]}"""
                # Sprint 61: Leid A/B sovereign separation
                if _tier == "vault":
                    from interfaces.config import VAULT_MAX_INPUT_TOKENS as _vmax
                    if _estimate_tokens(_msg) > _vmax:
                        return JSONResponse(status_code=413, content={
                            "error_code": "vault_input_too_large",
                            "detail": f"Skjal er of stórt fyrir Vault tier (max {_vmax} tokens). Skiptu í smærri hluta eða notaðu Almenna greiningu."})
                    logger.info(f"[ALVITUR] Sprint61 analyze_doc tier=vault calling leid_b")
                    async with _VAULT_SEMAPHORE_WS:
                        _summary, _model_used, _usage = await _call_leid_b(_msg)
                    if _summary is None:
                        return JSONResponse(status_code=503, content={
                            "error_code": "vault_local_unavailable",
                            "detail": "Trúnaðarþjónusta tímabundið ekki tiltæk. Local AI module er að ræsast — reyndu aftur eftir 1 mínútu."})
                    _pipeline_source_doc = f"local_vllm_{_model_used}"
                    _in_tok = _usage.get("prompt_tokens", 0); _out_tok = _usage.get("completion_tokens", 0)
                    logger.info(f"[ALVITUR] leid_b analyze_doc in={_in_tok} out={_out_tok}")
                    _domain_doc = "vault"
                else:
                    from datetime import datetime as _dt, timezone as _tz
                    from interfaces.specialist_prompts import classify as _classify, get_specialist_prompt as _get_prompt
                    _classify_text = (query.strip() if query and query.strip() else heildartexti[:500])
                    try:
                        _domain_doc = await _classify(_classify_text or "general")
                    except Exception:
                        _domain_doc = "general"
                    _domain_doc = _domain_doc or "general"
                    _now_str = _dt.now(_tz.utc).strftime("%Y-%m-%d %H:%M UTC")
                    # s68-hotfix: conditional honesty — document vs text-only mode
                    _has_document = bool(file and getattr(file, "filename", "") and file.filename.strip())
                    if _has_document:
                        _honesty_doc = (
                            "\n\nMikilvægt: Ef þú finnur ekki nógu nákvæmar upplýsingar í skjalinu "
                            "til að svara spurningunni, segjum notandanum það beint og bjóðum upp á "
                            "framhaldsspurningu. Búðu ALDREI til upplýsingar sem eru ekki í skjalinu."
                        )
                    else:
                        # s68-hotfix Part 3: tone-guide + thinking-suppress
                        _honesty_doc = (
                            "\n\nMikilvægt: Notandi hefur ekki hengt við skjal — þetta er almenn "
                            "spurning. Svaraðu út frá almennri þekkingu þinni á íslensku. "
                            "Ef þú veist ekki svarið með vissu, segðu það heiðarlega og bjóddu "
                            "framhaldsspurningu. Ekki búa til staðreyndir. "
                            "MIKILVÆGT UM FRAMSETNINGU: Ekki sýna hugsanaferli þitt — engar "
                            "<thinking>-tög, engin ‚bíð, ég þarf að“ setningar, engin internal-monologue. "
                            "Sendu aðeins hreint, prófessjónal svar beint. Engar emoji, "
                            "engin upphrópunarmerki, formlegt B2B-mál."
                        )
                    _system_prompt = _get_prompt(_domain_doc, _now_str) + _honesty_doc

                    # Sprint 70 Track D — RAG+ hook
                    print("DEBUG: RAG hook starting", flush=True)
                    try:
                        from core.rag_orchestrator import retrieve_legal_context, build_rag_injection
                        _rag = retrieve_legal_context(
                            query=((query or "").strip() or _msg or ""),
                            intent_domain=_domain_doc,
                            tier="general",
                            tenant_id="system",
                        )
                        if _rag.refusal:
                            return JSONResponse(content={
                                "success": True, "response": _rag.refusal,
                                "pipeline_source": "rag_refusal_vault",
                                "domain": _domain_doc, "zero_data": True,
                                "found": True, "status": "ready_for_analysis",
                                "citations": [], "quota_warning": None,
                            })
                        if _rag.used_retrieval:
                            _rag_injection = build_rag_injection(_rag.chunks)
                            _system_prompt = _system_prompt + _rag_injection
                            _pipeline_source_doc = f"rag_grounded_general"
                            _rag_doc_meta = {"used_retrieval": True, "chunks_count": len(_rag.chunks), "top_score": round(_rag.top_score,3), "low_confidence": _rag.top_score<0.65, "source_laws": []}
                            logger.info("[RAG] injected %d chunks into analyze_doc prompt", len(_rag.chunks))
                        elif _rag.fallback_to_gemini:
                            _system_prompt = _system_prompt + "[ATH: Engin lagatilvitnun fannst i corpus Alvitur. Svaraðu varlega.]"
                            _pipeline_source_doc = "rag_fallback_general"
                            _pipeline_source_doc = "rag_fallback_general"
                    except Exception as _rag_e:
                        logger.warning("[RAG] orchestrator villa (graceful degradation): %s", _rag_e)

                    _rag_doc_meta = {"used_retrieval": _rag.used_retrieval, "chunks_count": len(_rag.chunks), "top_score": round(_rag.top_score, 3), "low_confidence": _rag.top_score < 0.65, "source_laws": []} if "_rag" in dir() else {}
                    logger.info(f"[ALVITUR] Sprint61 analyze_doc tier=general calling leid_a domain={_domain_doc}")
                    _summary, _model_used, _usage = await _call_leid_a(_system_prompt, _msg)
                    if _summary is None:
                        # Sprint 62 Patch C: Leið A klikkaði → sovereign fallback á Leið B
                        logger.warning("[ALVITUR] Sprint62C analyze_doc: Leið A None, reyni fallback á Leið B")
                        try:
                            async with _VAULT_SEMAPHORE_WS:
                                _summary, _model_used, _usage = await _call_leid_b(_msg)
                        except Exception as _fe:
                            logger.error(f"[ALVITUR] Sprint62C analyze_doc fallback exc: {type(_fe).__name__}: {_fe}")
                            _summary = None
                        if _summary is None:
                            return JSONResponse(status_code=503, content={
                                "error_code": "both_pipelines_unavailable",
                                "detail": "Þjónusta tímabundið ekki aðgengileg. Reyndu eftir augnablik."})
                        _pipeline_source_doc = f"fallback_local_{_model_used}"
                        logger.info(f"[ALVITUR] Sprint62C analyze_doc fallback OK: model={_model_used}")
                    else:
                        _pipeline_source_doc = f"openrouter_{_model_used.split('/')[-1]}"
                    _in_tok = _usage.get("prompt_tokens", 0); _out_tok = _usage.get("completion_tokens", 0)
                    _cost = (_in_tok * 0.15 + _out_tok * 0.60) / 1_000_000 if "gpt-4o-mini" in _model_used else (_in_tok * 3.00 + _out_tok * 15.00) / 1_000_000 if "claude-sonnet" in _model_used else 0.0
                    logger.info("[ALVITUR] token_obs pipeline=analyze_doc model=%s tier=%s in=%d out=%d cost=%.6f", _model_used, _tier, _in_tok, _out_tok, _cost)
                    # Polish only for Leid A (never vault — sovereignty)
                    try:
                        # from interfaces.chat_routes import _polish as _polish_fn
                        _summary = await _polish_fn(_summary, _key)
                    except Exception as _pe:
                        logger.warning(f"[ALVITUR] polish failed (non-fatal): {_pe}")
        except Exception as _e:
            logger.error(f"[ALVITUR] analyze_doc pipeline exc: {type(_e).__name__}: {_e}")
            _summary = None
            _domain_doc = "general"
            _pipeline_source_doc = "error"
    response = {
        "success": True,
        "filename": file.filename,
        "sidur": sidur,
        "zero_data": True,
        "tier": _tier,
        "query": query or "",
        "response": _summary or result.get("response", ""),
        "domain": _domain_doc if "_domain_doc" in dir() else "general",
        "citations": result.get("citations", []),
        "found": result.get("found", False),
        "status": "ready_for_analysis",
        "quota_warning": _quota_warning,
        "rag_metadata": _rag_doc_meta,
        "pipeline_source": _pipeline_source_doc if "_pipeline_source_doc" in dir() else "unknown",
    }
    return JSONResponse(content=response)


@app.post("/api/chat")
async def chat_endpoint(request: Request):
    """Sprint 45: Production chat endpoint.
    Sprint 46 Phase 1: Quota check + CF-Connecting-IP fix.
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "Invalid JSON"})
    query = body.get("query", "").strip()
    # Hotfix: tier kemur annaðhvort úr JSON body eða X-Alvitur-Tier header.
    # Body hefur forgang — header er fallback fyrir eldri clients.
    _tier_body = body.get("tier", "").lower().strip()
    _tier_header = request.headers.get("X-Alvitur-Tier", "general").lower().strip()
    tier = _tier_body if _tier_body in ("general", "vault") else _tier_header
    if not query:
        return JSONResponse(status_code=422, content={"error_code": "empty_prompt"})

    # Sprint 46 Phase 1: IP quota check for /api/chat (was missing)
    import hashlib as _hl
    _master_key = os.environ.get("ALVITUR_MASTER_KEY_HASH", "")
    _req_key = request.headers.get("X-Master-Key", "")
    _is_admin = bool(_master_key and _req_key and _hl.sha256(_req_key.encode()).hexdigest() == _master_key)
    # CF-Connecting-IP has priority (real user IP behind Cloudflare)
    _client_ip = (
        request.headers.get("CF-Connecting-IP")
        or request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
        or (request.client.host if request.client else "unknown")
    )
    # ── Sprint 62: Beta check ──
    # Ef notandi skrifar beta-frasa → promotaður í 7 daga
    if _er_beta_fras(query):
        _promota_beta(_client_ip)
        logger.info(f"[BETA] {_client_ip} promotaður í beta-tier (7d)")
    _is_beta = _er_beta_ip(_client_ip)
    # ── /Sprint 62 Beta check ──
    _log_intent("chat", query, None, None, tier)

    _quota_count = _quota_tracker_chat.get(_client_ip, 0) + 1
    if not _is_admin and not _is_beta:
        _quota_tracker_chat[_client_ip] = _quota_count
    # Sprint 64 A1: samræma gate við tracker-update (bæta not _is_beta)
    if _quota_count > FREE_QUOTA and not _is_admin and not _is_beta:
        return JSONResponse(status_code=403, content={
            "success": False,
            "error": "Ókeypis prufutími er liðinn.",
            "error_code": "quota_exceeded",
            "upgrade_required": True,
            "upgrade_url": "/askrift",
            "message": "Þú hefur nýtt þr þr 5 ókeypis beiðnir. Uppfærðu til að halda áfram.",
        })

    from interfaces.chat_routes import handle_chat
    return await handle_chat(request, query, tier)


# ── Sprint 28: K6 — /oryggi trust page ───────────────────────────────────────
@app.get("/oryggi", response_class=HTMLResponse)
async def oryggi_page():
    """Sprint 29 T1 — Trust Center"""
    return HTMLResponse(content="""<!DOCTYPE html>
<html lang="is">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Öryggi | Alvitur</title>
  <meta name="description" content="Hvernig Alvitur meðhöndlar gögn, trúnað og skjalasendingar.">
  <meta property="og:title" content="Öryggi | Alvitur">
  <meta property="og:description" content="Ekkert vistast í trúnaðarham. Sjálfvirk gagnaeyðing. GDPR-samræmt.">
  <meta property="og:type" content="website">
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&display=swap" rel="stylesheet">
  <link rel="stylesheet" href="https://api.fontshare.com/v2/css?f[]=general-sans@400,500,600&display=swap">
  <style>
/* Alvitur.is Design System */
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
html { -webkit-font-smoothing: antialiased; scroll-behavior: smooth; scroll-padding-top: 5rem; }
:root {
  --font-display: 'General Sans', 'Helvetica Neue', sans-serif;
  --font-body: 'Inter', 'Helvetica Neue', sans-serif;
  --text-xs: clamp(0.75rem, 0.7rem + 0.25vw, 0.875rem);
  --text-sm: clamp(0.875rem, 0.8rem + 0.35vw, 1rem);
  --text-base: clamp(1rem, 0.95rem + 0.25vw, 1.125rem);
  --text-lg: clamp(1.125rem, 1rem + 0.75vw, 1.375rem);
  --text-hero: clamp(2.5rem, 1rem + 4vw, 4rem);
  --space-1: 0.25rem; --space-2: 0.5rem; --space-3: 0.75rem; --space-4: 1rem;
  --space-5: 1.25rem; --space-6: 1.5rem; --space-8: 2rem; --space-10: 2.5rem;
  --space-12: 3rem; --space-16: 4rem; --space-20: 5rem;
  --color-bg: #F5F4F0;
  --color-surface: #FFFFFF;
  --color-text: #1A1A1A;
  --color-text-muted: #6B6B6B;
  --color-text-faint: #A3A3A3;
  --color-accent: #0A6B6E;
  --color-accent-hover: #085456;
  --color-accent-light: rgba(10, 107, 110, 0.08);
  --color-accent-border: rgba(10, 107, 110, 0.25);
  --color-border: #E2E0DA;
  --color-border-light: #ECEAE5;
  --color-error: #B5364B;
  --color-success: #2D7A3E;
  --radius-sm: 0.375rem; --radius-md: 0.625rem; --radius-lg: 0.875rem; --radius-xl: 1rem;
  --shadow-card: 0 1px 3px rgba(26,26,26,0.04), 0 4px 16px rgba(26,26,26,0.06);
  --shadow-card-hover: 0 2px 8px rgba(26,26,26,0.06), 0 8px 24px rgba(26,26,26,0.08);
  --ease-out: cubic-bezier(0.16, 1, 0.3, 1);
  --transition-fast: 150ms var(--ease-out);
  --transition-normal: 250ms var(--ease-out);
  --content-max: 680px;
  --content-wide: 960px;
  --nav-height: 3.75rem;
}
body { min-height: 100dvh; line-height: 1.6; font-family: var(--font-body); font-size: var(--text-base); color: var(--color-text); background-color: var(--color-bg); }
h1,h2,h3,h4,h5,h6 { text-wrap: balance; line-height: 1.15; }
p,li { text-wrap: pretty; max-width: 72ch; }
button { cursor: pointer; background: none; border: none; font: inherit; color: inherit; }
/* NAV */
.nav { position: fixed; top: 0; left: 0; right: 0; z-index: 50; height: var(--nav-height); background: rgba(245,244,240,0.92); backdrop-filter: blur(12px); border-bottom: 1px solid var(--color-border-light); }
.nav__inner { max-width: var(--content-wide); margin: 0 auto; padding: 0 var(--space-6); height: 100%; display: flex; align-items: center; justify-content: space-between; }
.nav__logo { display: flex; align-items: center; gap: var(--space-2); text-decoration: none; color: var(--color-text); }
.nav__logo-text { font-family: var(--font-display); font-weight: 600; font-size: 1.25rem; letter-spacing: -0.02em; }
.nav__links { display: flex; align-items: center; gap: var(--space-6); }
.nav__link { font-size: var(--text-sm); font-weight: 500; color: var(--color-text-muted); text-decoration: none; }
.nav__link:hover { color: var(--color-accent); }
/* ORYGGI PAGE */
.oryggi-main { max-width: var(--content-max); margin: 0 auto; padding: calc(var(--nav-height) + var(--space-12)) var(--space-6) var(--space-16); }
.oryggi-main h1 { font-family: var(--font-display); font-size: clamp(1.75rem, 3vw, 2.25rem); font-weight: 600; letter-spacing: -0.02em; margin-bottom: var(--space-8); }
.oryggi-section { margin-bottom: var(--space-10); }
.oryggi-section h2 { font-family: var(--font-display); font-size: var(--text-lg); font-weight: 600; margin-bottom: var(--space-4); }
.oryggi-section p { font-size: var(--text-base); color: var(--color-text-muted); line-height: 1.7; margin-bottom: var(--space-4); }
.oryggi-cta { display: inline-flex; align-items: center; gap: var(--space-2); padding: var(--space-3) var(--space-6); background: var(--color-accent); color: #FFFFFF; font-size: var(--text-sm); font-weight: 600; text-decoration: none; border-radius: var(--radius-md); margin-top: var(--space-4); }
.oryggi-cta:hover { background: var(--color-accent-hover); }
.oryggi-trust { display: inline-flex; align-items: center; gap: var(--space-2); font-size: var(--text-xs); color: var(--color-text-faint); margin-top: var(--space-3); }
/* FOOTER */
.footer { padding: var(--space-8) var(--space-6); border-top: 1px solid var(--color-border-light); }
.footer__inner { max-width: var(--content-wide); margin: 0 auto; text-align: center; }
.footer__copy { font-size: var(--text-xs); color: var(--color-text-muted); margin-bottom: var(--space-2); }
.footer__copy a, .footer__links a { color: var(--color-text-muted); text-decoration: none; }
.footer__copy a:hover, .footer__links a:hover { color: var(--color-accent); }
.footer__links { font-size: var(--text-xs); color: var(--color-text-muted); margin-bottom: var(--space-2); }
.footer__links span { margin: 0 var(--space-2); }
.footer__eea { font-size: var(--text-xs); color: var(--color-text-faint); }
@media (max-width: 640px) {
  :root { --nav-height: 3.25rem; }
  .oryggi-main { padding-top: calc(var(--nav-height) + var(--space-8)); }
}
  </style>
</head>
<body>
  <nav class="nav" role="navigation" aria-label="Aðalvalmynd">
    <div class="nav__inner">
      <a href="/" class="nav__logo" aria-label="Alvitur forsíða">
        <svg class="nav__logo-mark" viewBox="0 0 28 28" fill="none" aria-hidden="true" width="28" height="28">
          <rect width="28" height="28" rx="6" fill="currentColor" opacity="0.1"/>
          <path d="M7 21L14 7L21 21M10.5 16h7" stroke="currentColor" stroke-width="2" stroke-linecap="round"/>
        </svg>
        <span class="nav__logo-text">Alvitur</span>
      </a>
      <div class="nav__links">
        <a href="/oryggi" class="nav__link">Öryggi</a>
        <a href="/" class="nav__link">&larr; Til baka</a>
      </div>
    </div>
  </nav>

  <main class="oryggi-main">
    <h1>Öryggi og gagnaforræði</h1>

    <section class="oryggi-section">
      <h2>Engin þjálfun á þínum gögnum</h2>
      <p>Hjá Alvitri teljum við að stofnanir og fyrirtæki eigi ekki að þurfa að velja á milli þess að nýta kraft öflugustu gervigreindar heims og þess að vernda viðkvæm gögn. Við höfum smíðað íslenskt vinnsluumhverfi og arkitektúr sem tryggir fullt gagnaforræði í hverju skrefi.</p>
      <p>Við nýtum eingöngu lokuð fyrirtækjaskil (Enterprise APIs) við stærstu mállíkön heims, þar sem gilda ströng skilyrði um gagnavernd. Við ábyrgjumst lagalega og tæknilega að hvorki skjöl né fyrirspurnir sem fara í gegnum Alvitur verði nokkurn tímann nýtt til að þjálfa, fínstilla eða bæta gervigreindarlíkön birgja okkar. Hugverkið þitt er varið.</p>
    </section>

    <section class="oryggi-section">
      <h2>Sjálfvirk gagnaeyðing í trúnaðarham</h2>
      <p>Fyrir viðkvæmustu upplýsingarnar krefst kerfið þess að notandi velji Trúnaðarvinnslu (Leið B). Undir þessu vinnslulagi eru skjöl lesin beint inn í vinnsluminni (RAM) á netþjónum okkar &mdash; þau snerta aldrei varanlegan harðan disk.</p>
      <p>Að greiningu lokinni er minnið hreinsað og gögnin eyðast sjálfkrafa. Engin varanleg afrit verða til. Það er ekki hægt að ná í gögn sem eru ekki lengur til.</p>
    </section>

    <section class="oryggi-section">
      <h2>Lögsaga og evrópskir innviðir</h2>
      <p>Allur vélbúnaður sem knýr gagnagátt Alviturs er hýstur í vottuðum gagnaverum innan Evrópska efnahagssvæðisins (EES). Kerfið lútur persónuverndarlögum (GDPR) að fullu.</p>
      <p>Með innbyggðri gagnaflokkun styður Alvitur við kröfur ISO 42001 og komandi löggjöf Evrópusambandsins (EU AI Act) um ábyrga og rekjanlega notkun gervigreindar.</p>
    </section>

    <section class="oryggi-section">
      <h2>Sannanlegur rekjanleiki</h2>
      <p>Til að styðja við örugga skjalavörslu og innri endurskoðun heldur kerfið utan um dulkóðaða atvikaskrá (Audit Trail). Við skráum að vinnsla á ákveðnu öryggisstigi fór fram, ásamt metadata og tímasetningu, en innihald gagnanna eða skjalsins sjálfs er aldrei varðveitt á Trúnaðarleiðinni.</p>
    </section>

    <div style="text-align: center; padding: var(--space-8) 0;">
      <a href="/" class="oryggi-cta">
        Hefja örugga greiningu
        <svg width="18" height="18" viewBox="0 0 18 18" fill="none" aria-hidden="true"><path d="M3.75 9h10.5M9.75 4.5L14.25 9l-4.5 4.5" stroke="currentColor" stroke-width="1.5" stroke-linecap="round" stroke-linejoin="round"/></svg>
      </a>
      <div class="oryggi-trust">
        <svg width="14" height="14" viewBox="0 0 14 14" fill="none" aria-hidden="true"><path d="M7 1L2 3.25v4C2 10.2 4.2 12.75 7 13.5c2.8-.75 5-3.3 5-6.25v-4L7 1z" stroke="currentColor" stroke-width="1.25" stroke-linejoin="round"/></svg>
        Styður við kröfur ISO 42001 og GDPR
      </div>
    </div>
  </main>

  <footer class="footer" role="contentinfo">
    <div class="footer__inner">
      <p class="footer__copy">&copy; 2026 Orkuskipti ehf &middot; <a href="mailto:info@alvitur.is">info@alvitur.is</a></p>
      <p class="footer__links"><a href="/personuvernd">Persónuverndarstefna</a><span>&middot;</span><a href="/skilmalar">Skilmálar</a></p>
      <p class="footer__eea">Gögn unnin innan EES</p>
    </div>
  </footer>
</body>
</html>""")
# ── Sprint 43b: /askrift pricing page ──────────────────────────────────────────
@app.get("/askrift", response_class=HTMLResponse)
async def askrift_page():
    """Sprint 43b — Pricing / subscription page."""
    try:
        from sprint43.pricing_page import render_pricing_page
        return HTMLResponse(content=render_pricing_page())
    except ImportError:
        return HTMLResponse(content="""<!DOCTYPE html>
<html lang="is"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Áskrift — Alvitur</title>
<style>body{font-family:system-ui,sans-serif;background:#0F1117;color:#E2E8F0;display:flex;justify-content:center;padding:3rem}
.c{max-width:600px;text-align:center}h1{margin-bottom:1rem}p{color:#94A3B8;margin-bottom:.5rem}
a{color:#6366F1}</style></head><body><div class="c">
<h1>Áskrift — Alvitur</h1>
<p>Brons: 990 kr/mán — 100 fyrirspurnir</p>
<p>Silfur: 4.990 kr/mán — 1.000 fyrirspurnir</p>
<p>Gull: 14.990 kr/mán — 5.000 fyrirspurnir</p>
<p style="margin-top:1.5rem"><a href="mailto:info@alvitur.is">Hafðu samband</a></p>
</div></body></html>""")


# ─────────────────────────────────────────────────────────────────────────────


@app.get("/mock-checkout/{plan}/{amount}/{user_id}", response_class=HTMLResponse)
async def checkout(plan: str, amount: int, user_id: str):
    return HTMLResponse(content="<h2 style='font-family:system-ui;color:#e5e5e5;background:#0a0a0a;padding:3rem;text-align:center'>Áskriftarlegar eru ekki opnar enn. Hafðu samband: info@alvitur.is</h2>", status_code=503)


@app.post("/api/webhook/mock_success", response_class=HTMLResponse)
async def mock_success(request: Request):
    # Samþykkja bæði form gögn og JSON
    content_type = request.headers.get("content-type", "")
    if "application/json" in content_type:
        data = await request.json()
    else:
        form = await request.form()
        data = dict(form)

    plan = data.get("plan", "unknown")
    user_id = data.get("user_id", "unknown")
    return HTMLResponse(content=build_success_page(plan, user_id))


# -- Sprint 66 pre-A: module-level PID lock file pointer (so refactor-safe) --
_pid_lock_fp = None
# -- /PID lock global --


if __name__ == "__main__":
    # -- Sprint 66 pre-A hotfix: single-instance PID lock --
    import fcntl as _fcntl_pl
    import sys as _sys_pl
    import os as _os_pl
    _PID_LOCK_PATH = _os_pl.getenv("ALVITUR_PID_LOCK", "/tmp/alvitur_web_server.lock")
    _pid_lock_fp = open(_PID_LOCK_PATH, "w", encoding="utf-8")
    try:
        _fcntl_pl.flock(_pid_lock_fp, _fcntl_pl.LOCK_EX | _fcntl_pl.LOCK_NB)
        _pid_lock_fp.write(str(_os_pl.getpid()))
        _pid_lock_fp.flush()
    except BlockingIOError:
        print(
            "[FATAL] web_server.py already running (lock: "
            + _PID_LOCK_PATH + ")",
            file=_sys_pl.stderr,
        )
        _sys_pl.exit(1)
    # -- /PID lock --
    uvicorn.run("web_server:app", host="0.0.0.0", port=int(__import__("os").environ.get("ALVITUR_PORT", "8000")))

@app.get("/alvitur-v2", response_class=HTMLResponse)
async def serve_v2():
    with open('index.html', 'r', encoding='utf-8') as f:
        return HTMLResponse(content=f.read())

@app.get("/mimir-demo", response_class=HTMLResponse)
async def serve_demo():
    with open('index.html', 'r', encoding='utf-8') as f:
        return HTMLResponse(content=f.read())


# ═══════════════════════════════════════════════════════════════════════
# Sprint 61 - Leid A/B helpers (sovereign separation)
# ═══════════════════════════════════════════════════════════════════════

def _estimate_tokens(text):
    return int(len((text or "").split()) * 1.3)
