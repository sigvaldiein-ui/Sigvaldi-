"""Build DOCX fixtures for C2 adapter tests (ASCII only)."""
from pathlib import Path
from docx import Document

OUT = Path("tests/fixtures")
OUT.mkdir(parents=True, exist_ok=True)

def build_sample():
    doc = Document()
    doc.add_heading("Sample Tables Doc", level=1)
    doc.add_paragraph("Paragraph before first table.")
    t1 = doc.add_table(rows=3, cols=3)
    t1.rows[0].cells[0].text = "Name"
    t1.rows[0].cells[1].text = "Age"
    t1.rows[0].cells[2].text = "City"
    t1.rows[1].cells[0].text = "Alice"
    t1.rows[1].cells[1].text = "30"
    t1.rows[1].cells[2].text = "Reykjavik"
    t1.rows[2].cells[0].text = "Bob"
    t1.rows[2].cells[1].text = "25"
    t1.rows[2].cells[2].text = "Akureyri"
    doc.add_paragraph("Between tables.")
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "Key"
    t2.rows[0].cells[1].text = "Value"
    t2.rows[1].cells[0].text = "alpha"
    t2.rows[1].cells[1].text = "42"
    path = OUT / "sample_tables.docx"
    doc.save(path)
    print(f"wrote {path} ({path.stat().st_size} bytes)")

def build_empty():
    doc = Document()
    doc.add_heading("No Tables Here", level=1)
    doc.add_paragraph("Just a paragraph, no tables at all.")
    path = OUT / "empty_notable.docx"
    doc.save(path)
    print(f"wrote {path} ({path.stat().st_size} bytes)")

if __name__ == "__main__":
    build_sample()
    build_empty()
    print("c2 fixtures ok")
